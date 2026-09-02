"""Schema-ordered DOCX assembly.

Sections with content are emitted in schema order. Empty sections are omitted.
Rating annotations require schema rating system + section has_rating_field + value.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

from backend.config import settings
from backend.domain.source_attribution import format_reference_source
from backend.models.report import ReportResult
from backend.models.schema import TemplateSchema
from backend.pii import scrubber as pii_scrubber
from backend.storage.photo_layout import SectionPhotoLayout, get_section_photo_layout

logger = logging.getLogger(__name__)

_DEFAULT_RATING_COLOURS = {"1": "00B050", "2": "ED7D31", "3": "C00000", "NI": "808080"}


def _apply_minimal_styles(doc) -> None:
    from docx.shared import Pt

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _rating_label(schema: TemplateSchema, value: str) -> str:
    fmt = schema.rating_system.format_template
    if fmt and "[VALUE]" in fmt:
        return fmt.replace("[VALUE]", value)
    name = schema.rating_system.name
    if name:
        return f"{name}: {value}"
    return str(value)


def _rating_colour(schema: TemplateSchema, value: str) -> str | None:
    for rv in schema.rating_system.values:
        if rv.value == value and rv.colour:
            return rv.colour.lstrip("#").upper()
    return _DEFAULT_RATING_COLOURS.get(str(value))


def _add_body_paragraph(doc, body: str):
    pii_scrubber.assert_no_pii(body, context="DOCX paragraph")
    return doc.add_paragraph(body)


def _add_reference_attribution(doc, gen) -> None:
    """Footnote listing uploaded past-report sources for this section.

    Gated on ``rag_sources``: only sections attributed during live generation
    carry these display strings. Reports rebuilt from provenance-only payloads
    intentionally omit the internal source footnote.
    """
    if not gen.rag_sources:
        return
    from docx.shared import Pt

    if gen.reference_sources:
        labels = [format_reference_source(src) for src in gen.reference_sources]
    else:
        labels = list(gen.rag_sources)
    text = "Source: " + "; ".join(labels) + "."
    pii_scrubber.assert_no_pii(text, context="DOCX source attribution")
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)


def _embed_section_photos(
    doc,
    paths: list[str | Path],
    layout: SectionPhotoLayout,
) -> None:
    """Insert section photos using placement learned from past reports."""
    from docx.shared import Inches, Pt, RGBColor

    clean = [Path(p) for p in paths if p and Path(p).is_file()]
    if not clean:
        return

    cap = doc.add_paragraph()
    cap_r = cap.add_run(layout.caption or "Photographs")
    cap_r.font.bold = True
    cap_r.font.size = Pt(9.5)
    cap_r.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    cap.paragraph_format.space_after = Pt(4)

    limit = settings.max_section_photos_per_section
    for photo_path in clean[:limit]:
        try:
            doc.add_picture(str(photo_path), width=Inches(layout.width_inches))
        except Exception as exc:  # noqa: BLE001
            logger.debug("Skipping corrupt photo %s: %s", photo_path, exc)
            continue

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(layout.spacing_after_pt)


def _build_content_body(
    doc,
    result: ReportResult,
    schema: TemplateSchema,
    *,
    rating_enabled: bool,
    photos_by_section: dict[str, list[str]],
) -> None:
    """Render content-mode sections grouped by topic (no RICS schema order)."""
    from docx.shared import RGBColor

    from backend.content_based import taxonomy

    ordered = sorted(
        result.sections,
        key=lambda s: (taxonomy.topic_order(s.topic_id), s.subtopic_id or ""),
    )
    last_topic = None
    for gen in ordered:
        if gen.status == "empty" or not (gen.text or "").strip():
            continue
        if gen.topic_id != last_topic:
            last_topic = gen.topic_id
            doc.add_heading(gen.topic_label or taxonomy.topic_label(gen.topic_id), level=1)

        heading = doc.add_heading(gen.title, level=2)
        if rating_enabled and gen.rating_value:
            label = _rating_label(schema, str(gen.rating_value))
            pii_scrubber.assert_no_pii(label, context="DOCX rating")
            run = heading.runs[0] if heading.runs else heading.add_run("")
            run.text = f"{run.text}  {label}" if heading.runs else f"  {label}"
            run.bold = True
            hexcol = _rating_colour(schema, str(gen.rating_value))
            if hexcol:
                run.font.color.rgb = RGBColor.from_string(hexcol)

        body = (gen.text or "").strip()
        if body:
            _add_body_paragraph(doc, body)

        section_paths = (
            photos_by_section.get(gen.section_id)
            or photos_by_section.get(gen.section_id.upper())
            or []
        )
        if section_paths:
            layout = get_section_photo_layout(result.tenant_id, gen.section_id)
            _embed_section_photos(doc, section_paths, layout)

        _add_reference_attribution(doc, gen)


def build_docx(
    result: ReportResult,
    schema: TemplateSchema,
    *,
    title: str = "Survey Report",
    include_footer: bool = False,
    template_docx_path: str | None = None,
    section_photo_paths: dict[str, list[str]] | None = None,
) -> bytes:
    from docx import Document
    from docx.shared import RGBColor

    full_text = (
        "\n".join(s.text for s in result.sections) + "\n" + result.unassigned_text
    )
    pii_scrubber.assert_no_pii(full_text, context="DOCX output")

    path = template_docx_path or settings.template_docx_path
    branded = Path(path) if path else None
    if branded and branded.is_file():
        doc = Document(str(branded))
    else:
        doc = Document()
        _apply_minimal_styles(doc)
        doc.add_heading(title, level=0)
        if schema.report_type:
            doc.add_paragraph(schema.report_type)
        if result.property_type:
            doc.add_paragraph(f"Property type: {result.property_type}")
        if result.tenure:
            doc.add_paragraph(f"Tenure: {result.tenure}")
        doc.add_paragraph(
            "[Address, client name, inspection date and surveyor details "
            "to be completed by the surveyor before issue.]"
        ).italic = True
        doc.add_page_break()

    rating_enabled = schema.rating_system.detected
    by_id = {s.section_id: s for s in result.sections}
    photos_by_section = section_photo_paths or {}

    # Content-based topic mode: sections carry a topic_id and are not part of the
    # RICS schema order, so render them grouped by topic instead.
    if any(getattr(s, "topic_id", "") for s in result.sections):
        _build_content_body(
            doc,
            result,
            schema,
            rating_enabled=rating_enabled,
            photos_by_section=photos_by_section,
        )
        schema_sections: list = []
    else:
        schema_sections = list(schema.ordered_sections())

    for sec in schema_sections:
        gen = by_id.get(sec.id)
        if gen is None or gen.status == "empty":
            continue

        heading_text = f"{sec.id}  {sec.title}" if sec.id else sec.title
        heading = doc.add_heading(heading_text, level=min(max(sec.level, 1), 4))

        if rating_enabled and sec.has_rating_field and gen.rating_value:
            label = _rating_label(schema, str(gen.rating_value))
            pii_scrubber.assert_no_pii(label, context="DOCX rating")
            if heading.runs:
                run = heading.runs[0]
                run.text = f"{run.text}  {label}"
                run.bold = True
                hexcol = _rating_colour(schema, str(gen.rating_value))
                if hexcol:
                    run.font.color.rgb = RGBColor.from_string(hexcol)
            else:
                run = heading.add_run(f"  {label}")
                run.bold = True
                hexcol = _rating_colour(schema, str(gen.rating_value))
                if hexcol:
                    run.font.color.rgb = RGBColor.from_string(hexcol)

        layout = get_section_photo_layout(result.tenant_id, sec.id)
        section_paths = (
            photos_by_section.get(sec.id) or photos_by_section.get(sec.id.upper()) or []
        )
        if layout.placement == "after_heading" and section_paths:
            _embed_section_photos(doc, section_paths, layout)

        if gen.status == "NO_RAG_MATCH":
            p = doc.add_paragraph()
            text = gen.text or (
                f'[No template paragraph matched for "{sec.title}". '
                f"Manual entry required.]"
            )
            pii_scrubber.assert_no_pii(text, context="DOCX paragraph")
            p.add_run(text).italic = True
        else:
            body = (gen.text or "").strip()
            if body:
                _add_body_paragraph(doc, body)

        if layout.placement == "after_body" and section_paths:
            _embed_section_photos(doc, section_paths, layout)

        _add_reference_attribution(doc, gen)

    unassigned = by_id.get("UNASSIGNED")
    if unassigned or result.unassigned_text.strip():
        doc.add_page_break()
        doc.add_heading("Unassigned Observations — Manual Review Required", level=2)
        doc.add_paragraph(
            "The following observations from the surveyor's notes could not be "
            "automatically assigned to any template section. Please review and draft manually."
        )
        obs = (
            unassigned.unmatched_observations
            if unassigned and unassigned.unmatched_observations
            else [result.unassigned_text]
        )
        for item in obs:
            if item.strip():
                bullet = f"• {item}"
                pii_scrubber.assert_no_pii(bullet, context="DOCX paragraph")
                doc.add_paragraph(bullet, style="List Bullet")

    show_footer = include_footer or settings.ai_transparency_footer_enabled
    if show_footer:
        doc.add_paragraph(
            "Generated from surveyor notes mapped onto the firm's master template."
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
