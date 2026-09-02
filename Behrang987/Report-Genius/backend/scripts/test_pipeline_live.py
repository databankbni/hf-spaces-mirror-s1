#!/usr/bin/env python3
"""Live integration smoke tests for Vision -> Mapping -> Grounding pipeline.

Run from the backend package root::

    cd backend
    python scripts/test_pipeline_live.py

Requires OPENAI_API_KEY in the environment (or .env at repo root).
"""

from __future__ import annotations

import io
import logging
import os
import sys
import tempfile
import traceback
from pathlib import Path
from typing import Any
from unittest.mock import patch

# ── Path bootstrap (backend/scripts -> repo root) ─────────────────────────────
REPO_ROOT = Path(__file__).resolve().parents[2]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))


# Prefer live key from the shell or repo .env before importing cached settings.
def _load_dotenv_file(path: Path) -> None:
    if not path.is_file():
        return
    for line in path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            continue
        key, _, value = stripped.partition("=")
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


_load_dotenv_file(REPO_ROOT / ".env")
_load_dotenv_file(REPO_ROOT / "backend" / ".env")

_env_key = (os.environ.get("OPENAI_API_KEY") or "").strip()
if _env_key:
    os.environ["OPENAI_API_KEY"] = _env_key

from backend.config import get_settings, settings  # noqa: E402

if _env_key:
    get_settings.cache_clear()
    settings.openai_api_key = _env_key

from docx import Document  # noqa: E402
from PIL import Image  # noqa: E402

from backend.domain import template_discoverer  # noqa: E402
from backend.ingest import pipeline as ingest  # noqa: E402
from backend.llm import openai_client  # noqa: E402
from backend.llm.openai_client import _parse_llm_json  # noqa: E402
from backend.pipeline import section_mapper  # noqa: E402
from backend.storage import photo_store  # noqa: E402

# ── Console helpers ─────────────────────────────────────────────────────────
GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
CYAN = "\033[96m"
RESET = "\033[0m"

_log_handler: logging.Handler | None = None
_captured_logs: list[str] = []


def _pass(msg: str) -> None:
    print(f"{GREEN}[PASS]{RESET} {msg}")


def _fail(msg: str) -> None:
    print(f"{RED}[FAIL]{RESET} {msg}")


def _info(msg: str) -> None:
    print(f"{CYAN}[INFO]{RESET} {msg}")


def _warn(msg: str) -> None:
    print(f"{YELLOW}[WARN]{RESET} {msg}")


def _section(title: str) -> None:
    print(f"\n{'=' * 72}\n{title}\n{'=' * 72}")


def _install_log_capture() -> None:
    global _log_handler, _captured_logs
    _captured_logs = []

    class _ListHandler(logging.Handler):
        def emit(self, record: logging.LogRecord) -> None:
            _captured_logs.append(record.getMessage())

    _log_handler = _ListHandler()
    _log_handler.setLevel(logging.INFO)
    logging.getLogger("backend.pipeline.section_mapper").addHandler(_log_handler)


def _remove_log_capture() -> None:
    global _log_handler
    if _log_handler is not None:
        logging.getLogger("backend.pipeline.section_mapper").removeHandler(_log_handler)
        _log_handler = None


def _log_contains(fragment: str) -> bool:
    return any(fragment in line for line in _captured_logs)


def _tiny_jpeg() -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (64, 64), color=(180, 180, 180)).save(buf, format="JPEG")
    return buf.getvalue()


def _write_fixture_docx(tmp: Path) -> tuple[Path, Path]:
    """Template + reference bundle with a known D2 baseline paragraph."""
    tpl = tmp / "template.docx"
    doc = Document()
    doc.add_paragraph("D2 Roof coverings")
    doc.save(str(tpl))

    past = tmp / "past.docx"
    doc2 = Document()
    doc2.add_paragraph("Section D2 - Roof Coverings")
    doc2.add_paragraph(
        "The roof covering comprised slate tiles laid to pitched timber rafters. "
        "Inspection was limited to ground level using binoculars."
    )
    doc2.save(str(past))
    return tpl, past


def _bootstrap_tenant(tmp: Path, tenant: str) -> str:
    """Ingest template + reference; return draft id with one selected D2 photo."""
    tpl, past = _write_fixture_docx(tmp)
    ingest.ingest_report_template(tenant, tpl)
    ingest.ingest_reference(tenant, past)
    template_discoverer.ensure_canonical_schema(tenant)

    draft = f"live-draft-{tenant}"
    row = photo_store.add_section_photo(
        tenant,
        draft,
        "D2",
        _tiny_jpeg(),
        content_type="image/jpeg",
    )
    photo_store.set_ai_selection(tenant, draft, "D2", [row.id])
    return draft


def _d2_section(result: Any) -> Any:
    for section in result.sections:
        if section.section_id.upper() == "D2":
            return section
    raise AssertionError("D2 section missing from report result")


def _guard_live_llm() -> bool:
    key = (os.environ.get("OPENAI_API_KEY") or settings.openai_api_key or "").strip()
    if not key:
        _fail("OPENAI_API_KEY is not set — export it or add it to .env at repo root.")
        return False
    settings.openai_api_key = key
    if not openai_client.is_available():
        _fail("openai_client.is_available() is False after loading OPENAI_API_KEY.")
        return False
    _info(f"OpenAI client ready (key prefix: {key[:7]}…).")
    return True


def test_case_a_pristine_passthrough(tmp: Path) -> bool:
    """Empty notes + zero photo observations → baseline passthrough, no auditor LLM."""
    _section("TEST CASE A — THE PRISTINE PASSTHROUGH")
    tenant = "live-passthrough"
    draft = _bootstrap_tenant(tmp, tenant)
    baseline_snippet = "binoculars"

    stabilize_calls = 0
    chat_text_calls = 0
    chat_json_async_calls = 0

    real_stabilize = section_mapper.stabilize_section
    real_chat_text = openai_client.chat_text
    real_chat_json_async = openai_client.chat_json_async

    async def _counting_stabilize(*args: Any, **kwargs: Any) -> Any:
        nonlocal stabilize_calls
        stabilize_calls += 1
        return await real_stabilize(*args, **kwargs)

    def _counting_chat_text(*args: Any, **kwargs: Any) -> str:
        nonlocal chat_text_calls
        chat_text_calls += 1
        return real_chat_text(*args, **kwargs)

    async def _counting_chat_json_async(*args: Any, **kwargs: Any) -> dict:
        nonlocal chat_json_async_calls
        chat_json_async_calls += 1
        return await real_chat_json_async(*args, **kwargs)

    try:
        _install_log_capture()
        with patch(
            "backend.pipeline.section_mapper.vision_observations_for_section",
            return_value=([], None),
        ):
            with patch.object(
                section_mapper, "stabilize_section", side_effect=_counting_stabilize
            ):
                with patch.object(
                    openai_client, "chat_text", side_effect=_counting_chat_text
                ):
                    with patch.object(
                        openai_client,
                        "chat_json_async",
                        side_effect=_counting_chat_json_async,
                    ):
                        result = section_mapper.generate_report(
                            tenant,
                            "",
                            interference_level="minimum",
                            report_draft_id=draft,
                            only_section_ids=["D2"],
                        )

        section = _d2_section(result)
        checks: list[tuple[bool, str]] = [
            (section.status == "OK", f"expected status OK, got {section.status!r}"),
            (
                baseline_snippet in section.text.lower(),
                "section text should retain retrieved baseline prose",
            ),
            (
                section.grounding_passed is True,
                f"expected grounding_passed=True, got {section.grounding_passed!r}",
            ),
            (
                stabilize_calls == 0,
                f"stabilize_section must not run (calls={stabilize_calls})",
            ),
            (
                chat_json_async_calls == 0,
                f"grounding auditor chat_json_async must not run (calls={chat_json_async_calls})",
            ),
            (
                _log_contains("baseline_passthrough_no_observations"),
                "expected baseline_passthrough_no_observations log line",
            ),
        ]

        ok = True
        for passed, detail in checks:
            if passed:
                _pass(detail)
            else:
                _fail(detail)
                ok = False

        if chat_text_calls:
            _warn(
                f"chat_text invoked {chat_text_calls} time(s) "
                "(notes expansion / mapping should be skipped for empty observations)."
            )
        else:
            _pass("no chat_text mapping calls during passthrough")

        return ok
    except Exception as exc:
        _fail(f"unexpected error: {exc}")
        traceback.print_exc()
        return False
    finally:
        _remove_log_capture()


def test_case_b_obscured_limitations_pass(tmp: Path) -> bool:
    """Limitations-only vision payload → grounding accepts defensive limitation prose."""
    _section("TEST CASE B — THE OBSCURED LIMITATIONS PASS")
    tenant = "live-limitations"
    draft = _bootstrap_tenant(tmp, tenant)
    limitation_phrase = "restricted view due to heavy insulation storage"
    vision_note = f"Photo limitations: {limitation_phrase}."

    try:
        with patch(
            "backend.pipeline.section_mapper.vision_observations_for_section",
            return_value=([], vision_note),
        ):
            result = section_mapper.generate_report(
                tenant,
                "",
                interference_level="maximum",
                report_draft_id=draft,
                only_section_ids=["D2"],
            )

        section = _d2_section(result)
        text_lower = section.text.lower()
        checks: list[tuple[bool, str]] = [
            (section.status == "OK", f"expected status OK, got {section.status!r}"),
            (
                section.grounding_passed is True,
                f"expected grounding_passed=True, got {section.grounding_passed!r}",
            ),
            (
                "insulation" in text_lower
                or "restricted" in text_lower
                or "view" in text_lower
                or limitation_phrase.lower() in text_lower,
                "section prose should reflect the photo limitation anchor",
            ),
            (
                vision_note.split("Photo limitations:", 1)[1].strip().lower()
                in text_lower
                or any(
                    token in text_lower
                    for token in ("insulation", "restricted", "view")
                ),
                "grounding should not strip limitation-defensive language as hallucination",
            ),
        ]

        ok = True
        for passed, detail in checks:
            if passed:
                _pass(detail)
            else:
                _fail(detail)
                ok = False

        _info(f"D2 output excerpt: {section.text[:220]}…")
        return ok
    except Exception as exc:
        _fail(f"unexpected error: {exc}")
        traceback.print_exc()
        return False


def test_case_c_json_fence_sanitization() -> bool:
    """_parse_llm_json strips markdown fences without JSONDecodeError."""
    _section("TEST CASE C — JSON FENCE SANITIZATION TRAP")
    samples = [
        '```json\n{"status": "ok"}\n```',
        '```\n{"status": "ok", "count": 1}\n```',
        '{"status": "bare"}',
    ]

    ok = True
    for index, raw in enumerate(samples, start=1):
        try:
            parsed = _parse_llm_json(raw)
            if not isinstance(parsed, dict):
                _fail(f"sample {index}: expected dict, got {type(parsed).__name__}")
                ok = False
                continue
            if parsed.get("status") != "ok" and index < 3:
                _fail(f"sample {index}: unexpected payload {parsed!r}")
                ok = False
                continue
            if index == 3 and parsed != {"status": "bare"}:
                _fail(f"sample {index}: unexpected payload {parsed!r}")
                ok = False
                continue
            _pass(f"sample {index}: fence-safe parse → {parsed!r}")
        except Exception as exc:
            _fail(f"sample {index}: raised {exc!r}")
            traceback.print_exc()
            ok = False

    # Optional live chat_json path: monkey-patch completion to return fenced JSON.
    if openai_client.is_available():
        try:
            fenced_payload = (
                '```json\n{"passed": true, "_audit_summary": "fence trap"}\n```'
            )

            def _fake_create(**kwargs: Any) -> Any:
                class _Msg:
                    content = fenced_payload

                class _Choice:
                    message = _Msg()

                class _Resp:
                    choices = [_Choice()]

                return _Resp()

            with patch.object(
                openai_client._get_client().chat.completions, "create", _fake_create
            ):
                live = openai_client.chat_json(
                    [{"role": "user", "content": "return json"}],
                    model=settings.grounding_model,
                )
            if live.get("passed") is True:
                _pass("chat_json live wrapper parsed fenced auditor JSON")
            else:
                _fail(f"chat_json live wrapper returned unexpected payload: {live!r}")
                ok = False
        except Exception as exc:
            _fail(f"chat_json fenced integration check failed: {exc}")
            traceback.print_exc()
            ok = False
    else:
        _warn("Skipping live chat_json fence trap (OpenAI unavailable).")

    return ok


def main() -> int:
    print(f"{CYAN}RICS pipeline live integration runner{RESET}")
    print(f"Repo root: {REPO_ROOT}")

    if not _guard_live_llm():
        return 1

    results: dict[str, bool] = {}

    with tempfile.TemporaryDirectory(prefix="rics-live-") as tmp_name:
        tmp = Path(tmp_name)
        settings.data_dir = str(tmp / "data")
        _info(f"Using isolated data_dir={settings.data_dir}")

        results["A_pristine_passthrough"] = test_case_a_pristine_passthrough(tmp)
        results["B_obscured_limitations"] = test_case_b_obscured_limitations_pass(tmp)
        results["C_json_fence_sanitization"] = test_case_c_json_fence_sanitization()

    _section("SUMMARY")
    failed = [name for name, ok in results.items() if not ok]
    for name, ok in results.items():
        if ok:
            _pass(name)
        else:
            _fail(name)

    if failed:
        print(f"\n{RED}{len(failed)} test block(s) failed:{RESET} {', '.join(failed)}")
        return 1

    print(f"\n{GREEN}All pipeline live checks passed.{RESET}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
