"""Extract raw note lines from uploaded field-notes files."""

from __future__ import annotations

import io
import re
from pathlib import Path

from backend.domain.notes.note_bullets import explode_note_lines


def _classify_notes_tier(text: str) -> tuple[int, float, str]:
    """Lightweight tier guess from note content (no LLM required)."""
    if len(text) < 200:
        return 0, 0.0, ""

    lower = text.lower()
    l3 = sum(
        1
        for kw in (
            "roof structure",
            "dpc",
            "damp proof",
            "chimney",
            "services",
            "drainage",
            "ground floor",
            "timber decay",
            "asbestos",
        )
        if kw in lower
    )
    l1 = sum(
        1 for kw in ("condition only", "no advice", "traffic light") if kw in lower
    )
    explicit = re.search(r"level\s*([123])", lower)
    if explicit:
        lvl = int(explicit.group(1))
        return lvl, 0.55, f"Mention of Level {lvl} in notes text."

    if l3 >= 3:
        return 3, 0.45, "Notes contain multiple Level 3-style inspection topics."
    if l1 >= 1 and l3 == 0:
        return 1, 0.35, "Notes suggest a condition-only (Level 1) style."
    if l3 >= 1:
        return (
            2,
            0.32,
            "Notes contain some detailed inspection topics (possible Level 2/3).",
        )
    return 0, 0.0, ""


def extract_lines_from_bytes(data: bytes, suffix: str) -> list[str]:
    """Parse docx/pdf/txt bytes into non-empty lines."""
    suffix = suffix.lower()
    lines: list[str] = []

    if suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(data))
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in lines:
                        lines.append(text)
    elif suffix == ".pdf":
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(data))
        for page in reader.pages:
            page_text = page.extract_text() or ""
            for line in page_text.splitlines():
                line = line.strip()
                if line:
                    lines.append(line)
    else:
        text = data.decode("utf-8", errors="replace")
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    return explode_note_lines(lines)


def extract_notes_from_upload(data: bytes, filename: str) -> dict:
    """Full /extract-notes response payload."""
    suffix = Path(filename or "").suffix.lower()
    if suffix not in {".docx", ".pdf", ".txt"}:
        raise ValueError(
            "Unsupported file type. Upload a .docx, .pdf, or .txt notes file."
        )

    lines = extract_lines_from_bytes(data, suffix)
    joined = "\n".join(lines)
    predicted_level, confidence, rationale = _classify_notes_tier(joined)
    if confidence and confidence < 0.30:
        predicted_level = 0
        rationale = (
            f"{rationale} (confidence {confidence:.0%} below threshold — treating as tier-unknown.)"
        ).strip()

    return {
        "filename": filename or "unknown",
        "lines": lines,
        "line_count": len(lines),
        "predicted_survey_level": predicted_level,
        "confidence": confidence,
        "rationale": rationale,
    }
