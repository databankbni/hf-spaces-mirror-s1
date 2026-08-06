"""Regression test suite for the SWMM Analysis MCP server.

Run:  SWMM_WORKER_PYTHON=$(which python) python test_regression.py [path/to/Kincora_Phase_2.inp]

Levels: (1) unit tests of deterministic screening logic; (2) an integration
run of the Kincora reference model from upload through report generation;
(3) DOCX content checks. Kincora expected values live ONLY here (as test
pins with documented tolerances), never in server code.

Tolerances: velocities ±0.05 m/s; depths ±0.005 m; volumes ±2 m³;
continuity ±0.02 percentage points; times ±120 s.
"""
from __future__ import annotations

import json
import math
import sys
from pathlib import Path

import pandas as pd

sys.path.insert(0, str(Path(__file__).parent))

PASS = FAIL = 0


def check(name: str, cond: bool, detail: str = "") -> None:
    global PASS, FAIL
    status = "PASS" if cond else "FAIL"
    if cond:
        PASS += 1
    else:
        FAIL += 1
    print(f"  {status}: {name}" + (f"  [{detail}]" if detail and not cond else ""))


# ===========================================================================
# 1. UNIT TESTS
# ===========================================================================

def unit_tests() -> None:
    print("\n== Unit tests: screening_logic ==")
    from screening_logic import (classify_velocity, continuity_disclosure,
                                 execution_integrity_assessment,
                                 effective_velocity_table,
                                 missing_information_register,
                                 resolve_legacy_solver_options,
                                 validate_solver_options)

    solver_errors = validate_solver_options({
        "FLOW_ROUTING": "DYNWAVE", "MAX_TRIALS": "0", "HEAD_TOLERANCE": "0"})
    check("legacy zero sentinels are not blocking", solver_errors == [])
    resolved = resolve_legacy_solver_options({
        "FLOW_ROUTING": "DYNWAVE", "FLOW_UNITS": "CMS",
        "MIN_SURFAREA": "0", "MAX_TRIALS": "0", "HEAD_TOLERANCE": "0"})
    effective = resolved["effective_options"]
    check("SI legacy MAX_TRIALS resolves to 8", effective["MAX_TRIALS"] == "8")
    check("SI legacy HEAD_TOLERANCE resolves to 0.0015", effective["HEAD_TOLERANCE"] == "0.0015")
    check("SI legacy MIN_SURFAREA resolves to 1.167", effective["MIN_SURFAREA"] == "1.167")
    check("three substitutions are audited", len(resolved["substitutions"]) == 3)
    check("negative values remain blocked", any("cannot be negative" in x for x in
          validate_solver_options({"FLOW_ROUTING": "DYNWAVE", "MAX_TRIALS": "-1"})))
    check("non-numeric values remain blocked", any("must be numeric" in x for x in
          validate_solver_options({"FLOW_ROUTING": "DYNWAVE", "HEAD_TOLERANCE": "bad"})))
    check("omitted dynamic-wave overrides use engine defaults",
          validate_solver_options({"FLOW_ROUTING": "DYNWAVE"}) == [])

    invalid = execution_integrity_assessment({
        "routing_steps": 7201, "not_converged_steps": 7201,
        "pct_not_converged": 100.0, "flow_error": 100.0,
        "runoff_error": -1.933})
    check("100% nonconvergence invalidates results", invalid["status"] == "invalid")
    check("invalid run blocks hydraulic conclusions",
          not invalid["results_usable"] and not invalid["hydraulic_conclusions_allowed"])
    valid = execution_integrity_assessment({
        "routing_steps": 12641, "not_converged_steps": 0,
        "pct_not_converged": 0.0, "flow_error": 0.022,
        "runoff_error": -1.933})
    check("corrected Kincora run remains usable", valid["results_usable"])

    check("advisory classification (3.5)", "Advisory" in classify_velocity(3.5))
    check("critical classification (4.5)", "Critical" in classify_velocity(4.5))
    check("boundary 3.0 is below advisory", "Below" in classify_velocity(3.0))
    check("boundary 4.0 is advisory not critical", "Advisory" in classify_velocity(4.0))
    check("None -> Not assessed", classify_velocity(None) == "Not assessed")
    check("NaN -> Not assessed", classify_velocity(float("nan")) == "Not assessed")

    md = {"runoff_error": -1.933, "flow_error": 0.022}
    lines = continuity_disclosure(md, 0.5, 1.0, has_pollutants=False)
    joined = " | ".join(lines)
    check("runoff sign preserved", "-1.933%" in joined)
    check("runoff exceeds 1% absolute warning", "⚠️" in joined and "1.933" in joined)
    check("routing reported separately, no warning", "+0.022%" in joined and joined.count("⚠️") == 1)
    check("quality N/A without pollutants", "Not applicable — no pollutants modelled" in joined)
    lines2 = continuity_disclosure({"runoff_error": -1.933, "flow_error": 0.022,
                                    "quality_error": 0.4}, 0.5, 1.0, has_pollutants=True)
    check("quality reported when pollutants modelled", any("+0.400%" in l for l in lines2))
    check("negative below-threshold not warned", "⚠️" not in " ".join(
        continuity_disclosure({"runoff_error": -0.4, "flow_error": 0.1}, 0.5, 1.0)))

    link_df = pd.DataFrame([
        {"Link ID": "L1", "Peak Velocity (m/s)": 1.186},
        {"Link ID": "L2", "Peak Velocity (m/s)": 5.6},
    ])
    recon = pd.DataFrame([
        {"Link ID": "L1", "RPT Peak Velocity": 2.63, "Overall Status": "Discrepancy"},
        {"Link ID": "L2", "RPT Peak Velocity": 5.61, "Overall Status": "OK"},
    ])
    eff = effective_velocity_table(link_df, recon)
    r1 = eff[eff["Link ID"] == "L1"].iloc[0]
    r2 = eff[eff["Link ID"] == "L2"].iloc[0]
    check("flagged link uses .rpt value", abs(r1["Screening Velocity (m/s)"] - 2.63) < 1e-6)
    check("flagged link source is .rpt", ".rpt" in r1["Evidence Source"])
    check("both values shown", r1["Worker Peak Velocity (m/s)"] == 1.186 and r1["RPT Peak Velocity (m/s)"] == 2.63)
    check("deltas computed", abs(r1["Delta (m/s)"] - (1.186 - 2.63)) < 1e-6 and r1["Delta (%)"] is not None)
    check("discrepancy does not create 3.0 exceedance", "Below advisory" in r1["Screening Classification"])
    check("classification-changed = No for L1", r1["Classification Changed by Reconciliation"] == "No")
    check("unflagged link uses worker value", abs(r2["Screening Velocity (m/s)"] - 5.6) < 1e-6
          and "worker" in r2["Evidence Source"])

    print("\n== Unit tests: storage & criteria gating ==")
    from calgary_rules import CalgaryCriteria, apply_storage_classification, classify_overland
    crit = CalgaryCriteria()
    stor = pd.DataFrame([
        {"Node ID": "Storage8a", "Type": "storage", "Time-Series Peak Depth (m)": 0.491, "Maximum Depth (m)": 0.5},
        {"Node ID": "SUB1", "Type": "storage", "Time-Series Peak Depth (m)": 0.349, "Maximum Depth (m)": 2.0},
    ])
    sc = apply_storage_classification(stor, crit, "m")
    joined = sc.to_string()
    check("no bare Pass in non-trap storage", not any(
        str(s).strip() == "Pass" for s in sc["Calgary Status"]))
    check("Not assessed language present for unclassified compliance",
          "Not assessed" in joined)
    margin_col = next(c for c in sc.columns if "Margin" in c)
    check("margin column renamed to Modelled", margin_col.startswith("Modelled Depth Margin"))
    m8a = float(sc[sc["Node ID"] == "Storage8a"][margin_col].iloc[0])
    check("Storage8a modelled margin ~= 0.009 m", abs(m8a - 0.009) < 0.005, f"got {m8a}")
    status, allowed = classify_overland(0.05, 1.0, crit.depth_velocity_curve)
    check("overland verdict uses screening wording", "Screens" in status)


# ===========================================================================
# 2. INTEGRATION TESTS (Kincora reference model)
# ===========================================================================

def integration_tests(inp_path: str) -> tuple[str, dict]:
    print("\n== Integration: Kincora pipeline ==")
    import tools
    inp_text = Path(inp_path).read_text(encoding="utf-8", errors="replace")
    up = tools.upload_model(inp_text, Path(inp_path).name)
    sid = up["session_id"]
    check("inventory: 16 conduits / 11 subs / 11 storage",
          up["element_counts"].get("CONDUITS") == 16
          and up["element_counts"].get("SUBCATCHMENTS") == 11
          and up["element_counts"].get("STORAGE") == 11, str(up["element_counts"]))

    run = tools.run_simulation(sid)
    check("simulation completed", run["simulation"] == "completed")
    check("model SHA-256 present", len(run.get("model_sha256", "")) == 64)
    check("runoff continuity ~ -1.933%", abs(run["runoff_continuity_error_pct"] - (-1.933)) <= 0.02)
    check("routing continuity ~ 0.022%", abs(run["flow_continuity_error_pct"] - 0.022) <= 0.02)
    check("reconciliation not claimed clean",
          "Review" in run["rpt_reconciliation"].get("verdict", ""))

    # High-velocity conduits (worker values ~= .rpt after Rev 23.2 velocity fix)
    links = {r["Link ID"]: r for r in tools.get_link_results(sid, limit=30)["rows"]}
    expected = {"1000": 5.606, "1001": 5.373, "1003": 3.864, "1002": 3.456, "1005": 3.251}
    for lid, ev in expected.items():
        got = links[lid]["Peak Velocity (m/s)"]
        check(f"velocity {lid} ~= {ev}", abs(got - ev) <= 0.05, f"got {got}")

    # Storage8a pins
    stor = {r["Node ID"]: r for r in tools.get_node_results(sid, node_type="storage", limit=30)["rows"]}
    s8 = stor["Storage8a"]
    check("Storage8a peak depth ~= 0.491 m", abs(s8["Peak Depth (m)"] - 0.491) <= 0.005)
    check("Storage8a utilization ~= 98.2%", abs(s8["Depth Ratio"] - 0.982) <= 0.01)
    ts = tools.get_timeseries(sid, "node", "Storage8a", "volume")
    check("Storage8a peak volume ~= 179.8 m3", abs(ts["peak"] - 179.8) <= 2.0, f"got {ts['peak']}")
    check("Storage8a peak time ~ 00:26:30",
          "00:26" in str(ts["time_of_peak"]) or "00:27" in str(ts["time_of_peak"]),
          str(ts["time_of_peak"]))

    # CB2A/CB2a case QA + 108(Spill) reconciliation
    review = tools.preliminary_design_review(sid)
    basis = " | ".join(str(r.get("deterministic_basis")) for r in review["findings"]["rows"])
    check("CB2A/CB2a case finding present", "CB2A" in basis and "CB2a" in basis and "case" in basis.lower())
    recon = tools.get_reconciliation(sid)
    fl = {r["Link ID"]: r for r in recon["flagged_links"]["rows"]}
    check("108(Spill) flagged", "108(Spill)" in fl)
    if "108(Spill)" in fl:
        r = fl["108(Spill)"]
        check("108(Spill) worker ~1.186 vs rpt ~2.63",
              abs(r["Worker Peak Velocity"] - 1.186) <= 0.05 and abs(r["RPT Peak Velocity"] - 2.63) <= 0.05,
              f"{r['Worker Peak Velocity']} / {r['RPT Peak Velocity']}")

    rep = tools.generate_report(sid, "Kincora Regression", client="Test")
    from sessions import STORE
    docx_path = STORE.get(sid).workdir / "outputs" / "Kincora_Regression_SWMR_Draft.docx"
    return str(docx_path), {"session": sid}


# ===========================================================================
# 3. DOCX CONTENT TESTS
# ===========================================================================

def docx_tests(docx_path: str) -> None:
    print("\n== DOCX content checks ==")
    from docx import Document
    doc = Document(docx_path)
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)
    heads = [t for t in texts if t and doc.paragraphs[texts.index(t)].style.name.startswith("Heading")]

    for h in ["1.0 INTRODUCTION", "4.9 Continuity and Result Reconciliation",
              "4.10 Missing-Information Register", "5.0 SUMMARY OF FINDINGS",
              "APPENDIX D - MODEL INPUT AND OUTPUT LISTINGS", "5.0a Prioritized Actions",
              "1.0a Model Identity and Execution Provenance"]:
        check(f"heading present: {h[:45]}", any(t.startswith(h) for t in texts))

    check("runoff warning disclosed with sign", "-1.933%" in joined and "⚠️" in joined)
    check("routing reported separately", "+0.022%" in joined)
    check("quality N/A disclosed", "Not applicable — no pollutants modelled" in joined)
    check("reconciliation not claimed clean", "not claimed to be fully clean" in joined)
    check("108(Spill) precedence statement present",
          "108(Spill)" in joined and ".rpt value governs" in joined)
    check("108(Spill) discrepancy does not change classification",
          "does not change the screening classification" in joined)
    all_cells = " | ".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    everything = joined + " | " + all_cells
    check("critical vs advisory language present",
          "CRITICAL screening exceedance" in everything and "Advisory screening exceedance" in everything)
    check("preliminary disclaimer present", "professional authentication are required" in joined)
    check("no calibration/validation claims",
          "calibrated" not in joined.lower() and "validated model" not in joined.lower())
    check("mandatory depth-velocity figure caption present",
          "Figure 9-1 - Alberta/Calgary Depth–Velocity Criteria" in joined)
    check("depth-velocity criterion table present",
          "Figure 9-1A - Permissible Depth and Velocity" in joined)
    check("automated model schematic caption present",
          "Figure 3-1 - Automated SWMM Model Schematic" in joined)

    # No unsupported storage Pass: scan tables for a bare "Pass" in Calgary Status
    bare_pass = False
    for tbl in doc.tables:
        headers = [c.text for c in tbl.rows[0].cells]
        if any("Calgary Status" in h for h in headers):
            idx = next(i for i, h in enumerate(headers) if "Calgary Status" in h)
            for row in tbl.rows[1:]:
                if row.cells[idx].text.strip() == "Pass":
                    bare_pass = True
    check("no bare 'Pass' in storage status table", not bare_pass)
    check("modelled-margin terminology used", "Modelled Depth Margin" in
          " ".join(c.text for t in doc.tables for r in t.rows for c in r.cells))
    check("SHA-256 in identity table", any(len(c.text.strip()) == 64 and all(
        ch in "0123456789abcdef" for ch in c.text.strip())
        for t in doc.tables[:6] for r in t.rows for c in r.cells))
    n_blank_runs = sum(1 for i in range(len(texts) - 3)
                       if all(not texts[i + j].strip() for j in range(4)))
    check("no 4+ consecutive blank paragraphs", n_blank_runs == 0, f"{n_blank_runs} runs")


if __name__ == "__main__":
    inp = sys.argv[1] if len(sys.argv) > 1 else "Kincora_Phase_2.inp"
    unit_tests()
    docx_path, _ = integration_tests(inp)
    docx_tests(docx_path)
    print(f"\n{'=' * 50}\nTOTAL: {PASS} passed, {FAIL} failed")
    sys.exit(1 if FAIL else 0)
