"""Municipal stormwater report generator for the SWMM6 GIS Space.

Creates an editable DOCX report and a ZIP package of supporting CSV/JSON data
from the deterministic simulation results already held in Streamlit session state.
"""
from __future__ import annotations

import io
import json
import re
import sqlite3
import tempfile
import zipfile
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Any, Mapping

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from calgary_rules import (
    CalgaryCriteria, apply_storage_classification, build_llm_report_context,
    build_minor_system_capacity_table, build_overland_compliance_table,
    criteria_register, infer_design_event,
)


def _depth_velocity_figure(
    compliance: pd.DataFrame,
    curve: tuple[tuple[float, float], ...],
    *,
    results_usable: bool,
    flow_unit: str,
) -> bytes:
    """Render the auditable Alberta/Calgary depth-velocity screening figure."""
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    from matplotlib.figure import Figure

    points = sorted((float(v), float(d)) for v, d in curve)
    fig = Figure(figsize=(7.2, 5.3), dpi=180, facecolor="white")
    FigureCanvasAgg(fig)
    ax = fig.add_subplot(111)

    # Extend the first permissible depth horizontally to zero velocity.  Do
    # not extrapolate beyond the final criterion point (3.0 m/s).
    boundary_x = [0.0] + [p[0] for p in points]
    boundary_y = [points[0][1]] + [p[1] for p in points]
    ax.axhspan(0.0, max(0.9, max(d for _, d in points) * 1.08),
               color="#f4cccc", alpha=0.42)
    ax.fill_between(boundary_x, 0.0, boundary_y, color="#d9ead3", alpha=0.72,
                    label="Acceptable screening region")
    ax.plot(boundary_x, boundary_y, color="#0057b8", linewidth=2.0,
            marker="s", markersize=4.5, markerfacecolor="white",
            label="Alberta/Calgary criterion")

    plotted = 0
    if results_usable and compliance is not None and not compliance.empty:
        dcol = next((c for c in compliance if c.startswith("Peak Depth (")), None)
        vcol = next((c for c in compliance if c.startswith("Peak Velocity (")), None)
        qcol = next((c for c in compliance if c.startswith("Peak Flow (")), None)
        if dcol and vcol:
            data = compliance.copy()
            data["_d"] = pd.to_numeric(data[dcol], errors="coerce")
            data["_v"] = pd.to_numeric(data[vcol], errors="coerce")
            data["_q"] = pd.to_numeric(data[qcol], errors="coerce") if qcol else float("nan")
            data = data.dropna(subset=["_d", "_v"])
            plotted = len(data)
            if plotted:
                if qcol and data["_q"].notna().any():
                    scatter = ax.scatter(data["_v"], data["_d"], c=data["_q"],
                                         cmap="viridis", marker="x", s=48,
                                         linewidths=1.4, zorder=4,
                                         label="Model results")
                    cbar = fig.colorbar(scatter, ax=ax, pad=0.02)
                    cbar.set_label(f"Peak flow ({flow_unit})", fontsize=8)
                else:
                    ax.scatter(data["_v"], data["_d"], color="black", marker="x",
                               s=48, linewidths=1.4, zorder=4, label="Model results")
                for _, row in data.iterrows():
                    label = str(row.get("Segment", "")).strip()
                    if label:
                        ax.annotate(label, (row["_v"], row["_d"]), xytext=(3, 3),
                                    textcoords="offset points", fontsize=6.5)

    ax.text(0.18, 0.10, "ACCEPTABLE", transform=ax.transAxes, color="#548235",
            fontsize=11, weight="bold", alpha=0.9)
    ax.text(0.68, 0.52, "NOT ACCEPTABLE", transform=ax.transAxes, color="#a61c00",
            fontsize=10, weight="bold", alpha=0.82)
    if not results_usable:
        ax.text(0.5, 0.02, "Model points withheld: execution-integrity gate failed",
                transform=ax.transAxes, ha="center", va="bottom", fontsize=8,
                color="#a61c00", bbox={"facecolor": "white", "edgecolor": "#a61c00", "pad": 3})
    elif plotted == 0:
        ax.text(0.5, 0.02, "No modeled overland-route depth/velocity pairs identified",
                transform=ax.transAxes, ha="center", va="bottom", fontsize=8,
                bbox={"facecolor": "white", "edgecolor": "#666666", "pad": 3})

    ax.set_xlim(0.0, 4.0)
    ax.set_ylim(0.0, max(0.9, max(d for _, d in points) * 1.08))
    ax.set_xlabel("Velocity (m/s)")
    ax.set_ylabel("Depth (m)")
    ax.set_title("Depth–Velocity Criteria for Overland Flow")
    ax.grid(True, color="#777777", linewidth=0.45, alpha=0.65)
    ax.legend(loc="upper right", fontsize=7.5, frameon=True)
    fig.tight_layout()
    out = io.BytesIO()
    fig.savefig(out, format="png", dpi=180, metadata={"Software": "SWMM Analysis MCP Server"})
    return out.getvalue()


def _model_schematic_figure(inp_sections: Mapping[str, list[list[str]]]) -> tuple[bytes, dict[str, Any]]:
    """Render a deterministic SWMM topology schematic from tokenized INP data."""
    from collections import defaultdict
    from matplotlib.backends.backend_agg import FigureCanvasAgg
    from matplotlib.figure import Figure
    from matplotlib.lines import Line2D

    node_sections = {"JUNCTIONS": "Junction", "DIVIDERS": "Junction",
                     "STORAGE": "Storage", "OUTFALLS": "Outfall"}
    nodes: dict[str, str] = {}
    for section, kind in node_sections.items():
        for row in inp_sections.get(section, []) or []:
            if row:
                nodes[str(row[0])] = kind
    link_sections = ("CONDUITS", "PUMPS", "ORIFICES", "WEIRS", "OUTLETS")
    links: list[tuple[str, str, str, str]] = []
    for section in link_sections:
        for row in inp_sections.get(section, []) or []:
            if len(row) >= 3:
                links.append((str(row[0]), str(row[1]), str(row[2]), section[:-1].title()))

    coords: dict[str, tuple[float, float]] = {}
    for row in inp_sections.get("COORDINATES", []) or []:
        if len(row) >= 3:
            try:
                coords[str(row[0])] = (float(row[1]), float(row[2]))
            except (TypeError, ValueError):
                pass

    # Coordinate-free models receive a stable downstream-rank layout.
    if not all(n in coords for n in nodes):
        outgoing: dict[str, list[str]] = defaultdict(list)
        incoming: dict[str, int] = defaultdict(int)
        for _, source, target, _ in links:
            outgoing[source].append(target)
            incoming[target] += 1
        rank = {n: 0 for n in nodes}
        queue = sorted(n for n in nodes if incoming[n] == 0)
        for source in queue:
            for target in sorted(outgoing[source]):
                rank[target] = max(rank.get(target, 0), rank[source] + 1)
                incoming[target] -= 1
                if incoming[target] == 0:
                    queue.append(target)
        groups: dict[int, list[str]] = defaultdict(list)
        for n in sorted(nodes):
            groups[rank.get(n, 0)].append(n)
        for r, names in groups.items():
            for i, name in enumerate(names):
                coords.setdefault(name, (float(r), float(-i)))
        coordinate_basis = "INP coordinates with deterministic topology fallback"
    else:
        coordinate_basis = "INP [COORDINATES]"

    vertices: dict[str, list[tuple[float, float]]] = defaultdict(list)
    for row in inp_sections.get("VERTICES", []) or []:
        if len(row) >= 3:
            try:
                vertices[str(row[0])].append((float(row[1]), float(row[2])))
            except (TypeError, ValueError):
                pass

    polygon_points: dict[str, list[tuple[float, float]]] = defaultdict(list)
    for row in inp_sections.get("POLYGONS", inp_sections.get("Polygons", [])) or []:
        if len(row) >= 3:
            try:
                polygon_points[str(row[0])].append((float(row[1]), float(row[2])))
            except (TypeError, ValueError):
                pass
    subcatchments = []
    for row in inp_sections.get("SUBCATCHMENTS", []) or []:
        if len(row) >= 3:
            sid, outlet = str(row[0]), str(row[2])
            pts = polygon_points.get(sid, [])
            if pts:
                pos = (sum(x for x, _ in pts) / len(pts), sum(y for _, y in pts) / len(pts))
            elif sid in coords:
                pos = coords[sid]
            else:
                ox, oy = coords.get(outlet, (0.0, 0.0))
                pos = (ox, oy + 1.0)
            subcatchments.append((sid, outlet, pos))

    fig = Figure(figsize=(10.5, 7.0), dpi=180, facecolor="white")
    FigureCanvasAgg(fig)
    ax = fig.add_subplot(111)
    link_colors = {"Conduit": "#4472c4", "Outlet": "#7030a0", "Pump": "#ed7d31",
                   "Orifice": "#a5a5a5", "Weir": "#ffc000"}
    for lid, source, target, kind in links:
        if source not in coords or target not in coords:
            continue
        path = [coords[source], *vertices.get(lid, []), coords[target]]
        color = link_colors.get(kind, "#4472c4")
        for a, b in zip(path[:-1], path[1:]):
            ax.plot([a[0], b[0]], [a[1], b[1]], color=color, linewidth=1.0, zorder=1)
        a, b = path[-2], path[-1]
        ax.annotate("", xy=b, xytext=a,
                    arrowprops={"arrowstyle": "-|>", "color": color, "lw": 1.0,
                                "shrinkA": 4, "shrinkB": 6}, zorder=2)
        mid = path[len(path) // 2]
        ax.annotate(lid, mid, xytext=(2, 2), textcoords="offset points",
                    fontsize=5.5, color="#333333")

    node_style = {
        "Junction": ("o", "#ffffff", "#1f4e79"),
        "Storage": ("s", "#fff2cc", "#bf9000"),
        "Outfall": ("v", "#f4cccc", "#990000"),
    }
    for name, kind in nodes.items():
        if name not in coords:
            continue
        x, y = coords[name]
        marker, face, edge = node_style[kind]
        ax.scatter([x], [y], marker=marker, s=48, facecolor=face, edgecolor=edge,
                   linewidth=1.1, zorder=4)
        ax.annotate(name, (x, y), xytext=(4, 4), textcoords="offset points",
                    fontsize=6.5, zorder=5)

    for sid, outlet, (x, y) in subcatchments:
        ax.scatter([x], [y], marker="D", s=42, facecolor="#e2f0d9",
                   edgecolor="#548235", linewidth=1.0, zorder=3)
        ax.annotate(f"SC {sid}", (x, y), xytext=(4, -8), textcoords="offset points",
                    fontsize=6.2, zorder=5)
        if outlet in coords:
            ax.annotate("", xy=coords[outlet], xytext=(x, y),
                        arrowprops={"arrowstyle": "->", "color": "#70ad47", "lw": 0.75,
                                    "linestyle": "--", "shrinkA": 4, "shrinkB": 6}, zorder=1)

    handles = [
        Line2D([0], [0], marker="D", color="none", markerfacecolor="#e2f0d9",
               markeredgecolor="#548235", label="Subcatchment"),
        *[Line2D([0], [0], marker=node_style[k][0], color="none",
                 markerfacecolor=node_style[k][1], markeredgecolor=node_style[k][2], label=k)
          for k in ("Junction", "Storage", "Outfall")],
        Line2D([0], [0], color="#4472c4", label="Hydraulic link / flow direction"),
        Line2D([0], [0], color="#70ad47", linestyle="--", label="Runoff routing"),
    ]
    ax.legend(handles=handles, loc="best", fontsize=7, frameon=True)
    ax.set_title("Automated SWMM Model Schematic", fontsize=13, weight="bold")
    ax.set_aspect("equal", adjustable="datalim")
    ax.axis("off")
    ax.margins(0.10)
    fig.tight_layout()
    out = io.BytesIO()
    fig.savefig(out, format="png", dpi=180, bbox_inches="tight",
                metadata={"Software": "SWMM Analysis MCP Server"})
    manifest = {
        "coordinate_basis": coordinate_basis,
        "node_count": len(nodes), "link_count": len(links),
        "subcatchment_count": len(subcatchments),
        "node_types": {kind: sum(1 for v in nodes.values() if v == kind)
                       for kind in sorted(set(nodes.values()))},
        "limitations": "Topology schematic only; confirm geometry, crossings, scale, and drawing reconciliation before issue.",
    }
    return out.getvalue(), manifest


@dataclass
class ReportCriteria:
    node_depth_ratio: float = 0.80
    minimum_freeboard: float = 0.50
    conduit_depth_ratio: float = 0.80
    velocity_threshold: float = 4.00      # critical screening threshold
    velocity_advisory: float = 3.00       # advisory screening threshold
    continuity_review: float = 0.50
    continuity_warning: float = 1.00
    suppress_empty_sections: bool = True
    major_link_ids: tuple[str, ...] = ()
    area_classification: dict[str, str] | None = None
    calgary_enabled: bool = True
    minor_release_rate_lps_ha: float | None = None
    trap_low_max_depth_m: float = 0.50
    entrance_grade_margin_m: float = 0.30
    conduit_capacity_warning_ratio: float = 0.95
    special_link_limits: dict[str, float] | None = None
    storage_classification: dict[str, str] | None = None
    outfall_classification: dict[str, str] | None = None
    checklist_overrides: dict[str, str] | None = None
    drawing_inventory: tuple[str, ...] = ()
    applicable_reports: tuple[str, ...] = ()


@dataclass
class ReportMetadata:
    project_name: str = "SWMM Project"
    client: str = "Not provided"
    consultant: str = "Not provided"
    consultant_file_no: str = "Not provided"
    subdivision_no: str = "Not provided"
    outline_plan_no: str = "Not provided"
    development_permit_no: str = "Not provided"
    design_storm: str = "Model design event"
    prepared_by: str = "Not provided"
    checked_by: str = "Not provided"
    report_date: str = ""
    municipality: str = "City of Calgary-style"
    contact_name: str = "Not provided"
    contact_email: str = "Not provided"
    legal_description: str = "Not provided"
    submission_status: str = "Preliminary"
    construction_drawing_no: str = "Not provided"
    development_agreement_no: str = "Not provided"

    def __post_init__(self) -> None:
        if not self.report_date:
            self.report_date = datetime.now().strftime("%B %d, %Y")


def _safe_name(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value).strip())
    return value.strip("_") or "SWMM_Project"


def _set_cell_text(cell, value: Any, bold: bool = False, size: float = 8.0) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("" if value is None else str(value))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_cell_shading(cell, fill: str = "D9EAD3") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def _format_value(value: Any) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    if isinstance(value, float):
        if abs(value) >= 1000:
            return f"{value:,.2f}"
        if abs(value) >= 10:
            return f"{value:,.3f}".rstrip("0").rstrip(".")
        return f"{value:,.4f}".rstrip("0").rstrip(".")
    return str(value)


def _add_df_table(
    doc: Document,
    title: str,
    df: pd.DataFrame,
    max_rows: int = 250,
    *,
    font_size: float = 7.4,
    landscape: bool = False,
) -> None:
    if landscape:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        _set_landscape(section)
    doc.add_heading(title, level=3)
    if df is None or df.empty:
        doc.add_paragraph("No applicable model records were identified.")
        return

    shown = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(shown.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _repeat_table_header(table.rows[0])
    for i, col in enumerate(shown.columns):
        _set_cell_text(table.rows[0].cells[i], col, bold=True, size=max(6.3, font_size - 0.2))
        _set_cell_shading(table.rows[0].cells[i])
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(shown.columns):
            _set_cell_text(cells[i], _format_value(row[col]), size=font_size)
    if len(df) > max_rows:
        doc.add_paragraph(f"Table truncated in the report at {max_rows:,} rows. Complete data are included in the ZIP package.")


@dataclass(frozen=True)
class UnitContext:
    flow_units: str
    system: str
    flow: str
    length: str
    area: str
    velocity: str
    rainfall: str
    storage: str


def _unit_context(flow_units: str) -> UnitContext:
    u = str(flow_units or "").upper()
    if u in {"CFS", "GPM", "MGD", "IMGD", "AFD"}:
        flow = {"CFS": "cfs", "GPM": "gpm", "MGD": "MGD", "IMGD": "IMGD", "AFD": "ac-ft/day"}.get(u, u)
        return UnitContext(u, "US Customary", flow, "ft", "ac", "ft/s", "in/hr", "ft³")
    flow = {"CMS": "m³/s", "LPS": "L/s", "MLD": "ML/day"}.get(u, u or "model units")
    return UnitContext(u or "UNKNOWN", "SI", flow, "m", "ha", "m/s", "mm/hr", "m³")


def _rename_native_columns(df: pd.DataFrame, units: UnitContext) -> pd.DataFrame:
    if df is None:
        return pd.DataFrame()
    out = df.copy()
    replacements = {
        "Area (ha)": f"Area ({units.area})",
        "Width (m)": f"Width ({units.length})",
        "Length (m)": f"Length ({units.length})",
        "Diameter (m)": f"Diameter ({units.length})",
        "Geom1 (m)": f"Geom1 ({units.length})",
        "Inlet Offset (m)": f"Inlet Offset ({units.length})",
        "Outlet Offset (m)": f"Outlet Offset ({units.length})",
        "Invert (m)": f"Invert ({units.length})",
        "Full Depth (m)": f"Full Depth ({units.length})",
        "Peak Depth (m)": f"Peak Depth ({units.length})",
        "Maximum HGL (m)": f"Maximum HGL ({units.length})",
        "Ground/Rim (m)": f"Ground/Rim ({units.length})",
        "Freeboard (m)": f"Freeboard ({units.length})",
        "Peak Flow (m³/s)": f"Peak Flow ({units.flow})",
        "Peak Runoff (m³/s)": f"Peak Runoff ({units.flow})",
        "Peak Flooding (m³/s)": f"Peak Flooding ({units.flow})",
        "Peak Inflow (m³/s)": f"Peak Inflow ({units.flow})",
        "Peak Velocity (m/s)": f"Peak Velocity ({units.velocity})",
        "Peak Rainfall (mm/h)": f"Peak Rainfall ({units.rainfall})",
    }
    if units.flow_units == "CFS":
        replacements["Total Runoff (m³)"] = "Runoff Volume (ft³)"
        replacements["Total Rainfall Depth"] = "Total Rainfall (in)"
        replacements["Runoff Depth"] = "Runoff Depth (in)"
    elif units.flow_units == "CMS":
        replacements["Total Runoff (m³)"] = "Runoff Volume (m³)"
        replacements["Total Rainfall Depth"] = "Total Rainfall (mm)"
        replacements["Runoff Depth"] = "Runoff Depth (mm)"
    elif units.flow_units == "LPS":
        replacements["Total Runoff (m³)"] = "Runoff Volume (L)"
        replacements["Total Rainfall Depth"] = "Total Rainfall (mm)"
        replacements["Runoff Depth"] = "Runoff Depth (mm)"
    else:
        replacements["Total Runoff (m³)"] = f"Integrated Runoff ({units.flow}-s)"
    return out.rename(columns={k: v for k, v in replacements.items() if k in out.columns})

def _inp_options(sections: dict[str, list[list[str]]]) -> dict[str, str]:
    result: dict[str, str] = {}
    for row in sections.get("OPTIONS", []):
        if len(row) >= 2:
            result[row[0].upper()] = " ".join(row[1:])
    return result


def _subcatchment_model_table(sections: dict[str, list[list[str]]], sub_summary: pd.DataFrame) -> pd.DataFrame:
    attrs: dict[str, dict[str, Any]] = {}
    for row in sections.get("SUBCATCHMENTS", []):
        if len(row) >= 7:
            try:
                attrs[row[0]] = {
                    "Rain Gage": row[1], "Outlet": row[2], "Area (ha)": float(row[3]),
                    "Impervious (%)": float(row[4]), "Width (m)": float(row[5]), "Slope (%)": float(row[6]),
                }
            except ValueError:
                continue
    for row in sections.get("SUBAREAS", []):
        if len(row) >= 7 and row[0] in attrs:
            try:
                attrs[row[0]].update({
                    "n Imperv.": float(row[1]), "n Perv.": float(row[2]),
                    "Dstore Imperv.": float(row[3]), "Dstore Perv.": float(row[4]),
                    "Zero Imperv. (%)": float(row[5]), "Routing": row[6],
                })
            except ValueError:
                pass
    model_df = pd.DataFrame([{"Sub ID": sid, **vals} for sid, vals in attrs.items()])
    if model_df.empty:
        return sub_summary.copy()
    if sub_summary is not None and not sub_summary.empty:
        # Model input values are authoritative for shared fields; avoid duplicate area/impervious columns.
        result_df = sub_summary.drop(columns=[c for c in ["Area (ha)", "% Impervious"] if c in sub_summary.columns], errors="ignore")
        return model_df.merge(result_df, on="Sub ID", how="left")
    return model_df


def _link_model_table(sections: dict[str, list[list[str]]], link_summary: pd.DataFrame) -> pd.DataFrame:
    records: dict[str, dict[str, Any]] = {}
    for section, link_type in (("CONDUITS", "conduit"), ("PUMPS", "pump"), ("ORIFICES", "orifice"), ("WEIRS", "weir"), ("OUTLETS", "outlet")):
        for row in sections.get(section, []):
            if len(row) < 3:
                continue
            rec = {"Link ID": row[0], "Model Type": link_type, "From Node": row[1], "To Node": row[2]}
            if section == "CONDUITS" and len(row) >= 9:
                try:
                    rec.update({
                        "Length (m)": float(row[3]), "Manning n": float(row[4]),
                        "Inlet Offset (m)": float(row[5]), "Outlet Offset (m)": float(row[6]),
                        "Initial Flow": float(row[7]), "Maximum Flow": float(row[8]),
                    })
                except ValueError:
                    pass
            records[row[0]] = rec
    for row in sections.get("XSECTIONS", []):
        if len(row) >= 3 and row[0] in records:
            records[row[0]]["Shape"] = row[1]
            try:
                records[row[0]]["Geom1 (m)"] = float(row[2])
                if len(row) > 3: records[row[0]]["Geom2"] = float(row[3])
                if len(row) > 4: records[row[0]]["Geom3"] = float(row[4])
                if len(row) > 5: records[row[0]]["Geom4"] = float(row[5])
                if len(row) > 6: records[row[0]]["Barrels"] = int(float(row[6]))
            except ValueError:
                pass
    model_df = pd.DataFrame(records.values())
    if model_df.empty:
        return link_summary.copy()
    if link_summary is not None and not link_summary.empty:
        drop_overlap = [c for c in ["From Node", "To Node", "Length (m)"] if c in link_summary.columns]
        result_df = link_summary.drop(columns=drop_overlap, errors="ignore")
        return model_df.merge(result_df, on="Link ID", how="left")
    return model_df


def _node_hgl_table(node_summary: pd.DataFrame) -> pd.DataFrame:
    if node_summary is None or node_summary.empty:
        return pd.DataFrame()
    df = node_summary.copy()
    df["Maximum HGL (m)"] = df.get("Invert (m)", 0) + df.get("Peak Depth (m)", 0)
    df["Ground/Rim (m)"] = df.get("Invert (m)", 0) + df.get("Full Depth (m)", 0)
    df["Freeboard (m)"] = df["Ground/Rim (m)"] - df["Maximum HGL (m)"]
    cols = ["Node ID", "Type", "Invert (m)", "Ground/Rim (m)", "Maximum HGL (m)", "Peak Depth (m)", "Freeboard (m)", "Peak Flooding (m³/s)", "Status"]
    return df[[c for c in cols if c in df.columns]]



def _parse_time_seconds(value: str | None) -> float:
    """Parse SWMM HH:MM:SS or numeric-second time values."""
    if value is None:
        return 0.0
    text = str(value).strip()
    if not text:
        return 0.0
    try:
        return float(text)
    except ValueError:
        pass
    parts = text.split(":")
    try:
        nums = [float(x) for x in parts]
    except ValueError:
        return 0.0
    if len(nums) == 3:
        return nums[0] * 3600.0 + nums[1] * 60.0 + nums[2]
    if len(nums) == 2:
        return nums[0] * 60.0 + nums[1]
    return 0.0


def _enhance_subcatchment_results(df: pd.DataFrame, units: UnitContext, report_step_s: float) -> pd.DataFrame:
    """Add defensible event coefficients and integrated runoff volumes in native units."""
    if df is None or df.empty:
        return pd.DataFrame()
    out = df.copy()
    out = out.loc[:, ~out.columns.duplicated()].copy()
    area_col = next((c for c in out.columns if c.startswith("Area (")), None)
    q_col = next((c for c in out.columns if c.startswith("Peak Runoff (")), None)
    rain_col = next((c for c in out.columns if c.startswith("Peak Rainfall (")), None)
    sum_col = next((c for c in out.columns if c.startswith("Runoff Sum (")), None)

    if area_col and q_col and rain_col:
        area = pd.to_numeric(out[area_col], errors="coerce")
        q = pd.to_numeric(out[q_col], errors="coerce")
        intensity = pd.to_numeric(out[rain_col], errors="coerce")
        denom = pd.Series(float("nan"), index=out.index)
        if units.flow_units == "CFS":
            denom = 1.008 * intensity * area  # Q(cfs)=1.008*C*i(in/hr)*A(ac)
        elif units.flow_units == "CMS":
            denom = 0.0027777778 * intensity * area  # Q(m3/s)=0.0027778*C*i(mm/hr)*A(ha)
        elif units.flow_units == "LPS":
            denom = 2.7777778 * intensity * area  # Q(L/s)=2.7778*C*i(mm/hr)*A(ha)
        coeff = q / denom.where(denom > 0)
        out["Peak-Flow Runoff Coefficient"] = coeff.where((coeff >= 0) & (coeff <= 1.5))
        out.drop(columns=["Runoff Coefficient"], errors="ignore", inplace=True)

    if sum_col and report_step_s > 0:
        runoff_sum = pd.to_numeric(out[sum_col], errors="coerce").fillna(0.0)
        if units.flow_units == "CFS":
            volume_ft3 = runoff_sum * report_step_s
            out["Runoff Volume (ft³)"] = volume_ft3
            out["Runoff Volume (ac-ft)"] = volume_ft3 / 43560.0
            out["Runoff Volume (MG)"] = volume_ft3 * 7.48051948 / 1_000_000.0
            if area_col:
                area = pd.to_numeric(out[area_col], errors="coerce")
                out["Runoff Depth (in)"] = (volume_ft3 / (area * 43560.0)) * 12.0
        elif units.flow_units == "CMS":
            volume_m3 = runoff_sum * report_step_s
            out["Runoff Volume (m³)"] = volume_m3
            if area_col:
                area = pd.to_numeric(out[area_col], errors="coerce")
                out["Runoff Depth (mm)"] = volume_m3 / (area * 10.0)
        elif units.flow_units == "LPS":
            volume_m3 = runoff_sum * report_step_s / 1000.0
            out["Runoff Volume (m³)"] = volume_m3
            if area_col:
                area = pd.to_numeric(out[area_col], errors="coerce")
                out["Runoff Depth (mm)"] = volume_m3 / (area * 10.0)
        elif units.flow_units in {"GPM", "MGD", "IMGD", "MLD", "AFD"}:
            out[f"Integrated Runoff ({units.flow}-s)"] = runoff_sum * report_step_s
        out.drop(columns=[sum_col], errors="ignore", inplace=True)

    # Add convenient secondary volume units when the application already supplied an integrated native volume.
    if units.flow_units == "CFS" and "Runoff Volume (ft³)" in out.columns:
        volume_ft3 = pd.to_numeric(out["Runoff Volume (ft³)"], errors="coerce")
        out["Runoff Volume (ac-ft)"] = volume_ft3 / 43560.0
        out["Runoff Volume (MG)"] = volume_ft3 * 7.48051948 / 1_000_000.0
    elif units.flow_units == "LPS" and "Runoff Volume (L)" in out.columns:
        out["Runoff Volume (m³)"] = pd.to_numeric(out["Runoff Volume (L)"], errors="coerce") / 1000.0
    return out


def _populate_outfall_flows(node_df: pd.DataFrame, link_df: pd.DataFrame, units: UnitContext) -> pd.DataFrame:
    """Populate outfall peak inflow from incoming links when node output is absent or zero."""
    if node_df is None or node_df.empty:
        return pd.DataFrame()
    out = node_df.copy()
    inflow_col = f"Peak Inflow ({units.flow})"
    flow_col = f"Peak Flow ({units.flow})"
    if inflow_col not in out.columns:
        out[inflow_col] = 0.0
    if link_df is None or link_df.empty or "To Node" not in link_df or flow_col not in link_df:
        return out
    incoming = link_df.groupby("To Node")[flow_col].max()
    mask = out.get("Type", "").astype(str).str.lower().eq("outfall")
    for idx in out.index[mask]:
        node_id = out.at[idx, "Node ID"]
        current = pd.to_numeric(pd.Series([out.at[idx, inflow_col]]), errors="coerce").iloc[0]
        fallback = incoming.get(node_id)
        if (pd.isna(current) or float(current) <= 0.0) and fallback is not None and not pd.isna(fallback):
            out.at[idx, inflow_col] = float(fallback)
    return out


def _split_link_appendix(link_table: pd.DataFrame, units: UnitContext) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    geom_cols = ["Link ID", "From Node", "To Node", "Model Type", "Shape", f"Length ({units.length})",
                 f"Geom1 ({units.length})", "Geom2", "Geom3", "Geom4", "Barrels"]
    params_cols = ["Link ID", "Manning n", f"Inlet Offset ({units.length})", f"Outlet Offset ({units.length})",
                   "Initial Flow", "Maximum Flow"]
    results_cols = ["Link ID", f"Peak Flow ({units.flow})", f"Peak Depth ({units.length})", "Depth Ratio",
                    f"Peak Velocity ({units.velocity})", "Status"]
    return tuple(link_table[[c for c in cols if c in link_table.columns]].copy() for cols in (geom_cols, params_cols, results_cols))


def _split_subcatchment_appendix(sub_table: pd.DataFrame, units: UnitContext) -> tuple[pd.DataFrame, pd.DataFrame]:
    model_cols = ["Sub ID", "Rain Gage", "Outlet", f"Area ({units.area})", "Impervious (%)", f"Width ({units.length})",
                  "Slope (%)", "n Imperv.", "n Perv.", "Dstore Imperv.", "Dstore Perv.", "Zero Imperv. (%)", "Routing", "Connected To"]
    result_cols = ["Sub ID", f"Peak Runoff ({units.flow})", "Peak-Flow Runoff Coefficient",
                   f"Peak Rainfall ({units.rainfall})", "Runoff Volume (ft³)", "Runoff Volume (ac-ft)",
                   "Runoff Volume (MG)", "Runoff Volume (m³)", "Runoff Depth (in)", "Runoff Depth (mm)"]
    return (sub_table[[c for c in model_cols if c in sub_table.columns]].copy(),
            sub_table[[c for c in result_cols if c in sub_table.columns]].copy())



def _read_sqlite_tables(result_db_bytes: bytes | None, table_names: list[str]) -> dict[str, pd.DataFrame]:
    """Read selected result tables from the in-memory SQLite export package."""
    if not result_db_bytes:
        return {}
    tables: dict[str, pd.DataFrame] = {}
    tmp_name = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".sqlite", delete=False) as tmp:
            tmp.write(result_db_bytes)
            tmp_name = tmp.name
        with sqlite3.connect(tmp_name) as con:
            available = {r[0] for r in con.execute("SELECT name FROM sqlite_master WHERE type='table'")}
            for name in table_names:
                if name in available:
                    tables[name] = pd.read_sql_query(f'SELECT * FROM "{name}"', con)
    except Exception:
        return tables
    finally:
        if tmp_name:
            try:
                Path(tmp_name).unlink(missing_ok=True)
            except Exception:
                pass
    return tables


def _control_table(sections: dict[str, list[list[str]]], link_table: pd.DataFrame, units: UnitContext) -> pd.DataFrame:
    """Create a control-specific schedule for orifices, weirs and outlets."""
    peak_col = f"Peak Flow ({units.flow})"
    peak_lookup = {}
    if link_table is not None and not link_table.empty and peak_col in link_table.columns:
        peak_lookup = link_table.set_index("Link ID")[peak_col].to_dict()
    rows: list[dict[str, Any]] = []
    for row in sections.get("ORIFICES", []):
        if len(row) >= 6:
            rows.append({
                "Control ID": row[0], "Control Type": "Orifice", "From Node": row[1], "To Node": row[2],
                "Subtype": row[3], f"Offset ({units.length})": _num(row[4]), "Discharge Coefficient": _num(row[5]),
                "Flap Gate": row[6] if len(row) > 6 else "", "Opening/Closing Time": _num(row[7]) if len(row) > 7 else None,
                peak_col: peak_lookup.get(row[0]),
            })
    for row in sections.get("WEIRS", []):
        if len(row) >= 6:
            rows.append({
                "Control ID": row[0], "Control Type": "Weir", "From Node": row[1], "To Node": row[2],
                "Subtype": row[3], f"Crest Height ({units.length})": _num(row[4]), "Discharge Coefficient": _num(row[5]),
                "Flap Gate": row[6] if len(row) > 6 else "", "End Contractions": _num(row[7]) if len(row) > 7 else None,
                "End Coefficient": _num(row[8]) if len(row) > 8 else None, "Surcharge Allowed": row[9] if len(row) > 9 else "",
                peak_col: peak_lookup.get(row[0]),
            })
    for row in sections.get("OUTLETS", []):
        if len(row) >= 6:
            rows.append({
                "Control ID": row[0], "Control Type": "Outlet", "From Node": row[1], "To Node": row[2],
                f"Offset ({units.length})": _num(row[3]), "Rating Type": row[4], "Rating Curve/Parameters": " ".join(row[5:]),
                peak_col: peak_lookup.get(row[0]),
            })
    return pd.DataFrame(rows)


def _num(value: Any) -> float | None:
    try:
        return float(value)
    except Exception:
        return None


def _control_reconciliation(control_df: pd.DataFrame, link_table: pd.DataFrame, units: UnitContext) -> pd.DataFrame:
    """Compare parallel control peaks with the downstream conveyance peak as a screening check."""
    peak_col = f"Peak Flow ({units.flow})"
    if control_df is None or control_df.empty or link_table is None or link_table.empty or peak_col not in link_table.columns:
        return pd.DataFrame()
    rows = []
    for to_node, grp in control_df.groupby("To Node", dropna=False):
        control_sum = pd.to_numeric(grp[peak_col], errors="coerce").fillna(0).sum()
        downstream = link_table[(link_table.get("From Node") == to_node) & (link_table.get("Model Type", link_table.get("Type", "")).astype(str).str.lower() == "conduit")]
        downstream_peak = pd.to_numeric(downstream[peak_col], errors="coerce").max() if not downstream.empty else float("nan")
        diff = downstream_peak - control_sum if pd.notna(downstream_peak) else float("nan")
        pct = abs(diff) / max(abs(downstream_peak), 1e-12) * 100 if pd.notna(diff) else float("nan")
        rows.append({
            "Receiving Node": to_node, "Controls": ", ".join(grp["Control ID"].astype(str)),
            f"Sum of Control Peaks ({units.flow})": control_sum,
            "Downstream Link": ", ".join(downstream["Link ID"].astype(str)) if not downstream.empty else "Not identified",
            f"Downstream Peak ({units.flow})": downstream_peak,
            f"Difference ({units.flow})": diff, "Difference (%)": pct,
            "Check": "Consistent" if pd.notna(pct) and pct <= 2.0 else "Review timing / routing",
        })
    return pd.DataFrame(rows)


def _storage_table(sections: dict[str, list[list[str]]], node_report: pd.DataFrame, db_tables: dict[str, pd.DataFrame], units: UnitContext) -> pd.DataFrame:
    """Create storage-specific input/result table using node time series when available."""
    node_ts = db_tables.get("node_timeseries", pd.DataFrame())
    node_lookup = node_report.set_index("Node ID").to_dict("index") if node_report is not None and not node_report.empty else {}
    rows = []
    for r in sections.get("STORAGE", []):
        if len(r) < 5:
            continue
        sid = r[0]; invert = _num(r[1]); max_depth = _num(r[2]); init_depth = _num(r[3]); shape = r[4]
        rec: dict[str, Any] = {"Storage ID": sid, f"Invert ({units.length})": invert, f"Maximum Depth ({units.length})": max_depth,
                               f"Initial Depth ({units.length})": init_depth, "Shape/Curve": shape,
                               "Shape Parameters": " ".join(r[5:])}
        nrec = node_lookup.get(sid, {})
        rec[f"Reported Peak Depth ({units.length})"] = nrec.get(f"Peak Depth ({units.length})")
        rec[f"Maximum HGL ({units.length})"] = nrec.get(f"Invert ({units.length})", invert) + (nrec.get(f"Peak Depth ({units.length})") or 0) if invert is not None else None
        if not node_ts.empty and "node_id" in node_ts.columns:
            g = node_ts[node_ts["node_id"].astype(str) == sid].copy()
            if not g.empty:
                for src, label in (("depth", f"Time-Series Peak Depth ({units.length})"), ("volume", f"Maximum Stored Volume ({units.storage})"),
                                   ("inflow", f"Peak Inflow ({units.flow})"), ("outflow", f"Peak Outflow ({units.flow})")):
                    if src in g:
                        rec[label] = pd.to_numeric(g[src], errors="coerce").max()
                if "volume" in g and "timestamp" in g:
                    vals = pd.to_numeric(g["volume"], errors="coerce")
                    if vals.notna().any(): rec["Time of Maximum Storage"] = g.loc[vals.idxmax(), "timestamp"]
        peak_depth = rec.get(f"Time-Series Peak Depth ({units.length})", rec.get(f"Reported Peak Depth ({units.length})"))
        rec["Depth Utilization (%)"] = 100 * float(peak_depth) / float(max_depth) if peak_depth is not None and max_depth and max_depth > 0 else None
        rec["Status"] = "Review zero storage response" if (peak_depth is not None and float(peak_depth) == 0 and (rec.get(f"Peak Outflow ({units.flow})") or 0) > 0) else "OK"
        rows.append(rec)
    return pd.DataFrame(rows)


def _area_classification(sub_table: pd.DataFrame, units: UnitContext, overrides: dict[str, str] | None = None) -> pd.DataFrame:
    """Separate likely proposed, pre-development and external/comparison catchments by naming convention."""
    if sub_table is None or sub_table.empty:
        return pd.DataFrame()
    area_col = f"Area ({units.area})"
    if area_col not in sub_table:
        return pd.DataFrame()
    overrides = {str(k): str(v) for k, v in (overrides or {}).items()}
    def classify(name: str) -> str:
        if str(name) in overrides:
            return overrides[str(name)]
        n = str(name).lower().replace("_", " ")
        if any(k in n for k in ("predevelop", "pre develop", "pre-development", "existing", "predev")):
            return "Pre-development / comparison"
        if any(k in n for k in ("external", "offsite", "off-site", "upstream")):
            return "External area"
        return "Proposed / modelled development"
    d = sub_table[["Sub ID", area_col]].copy(); d["Area Category"] = d["Sub ID"].map(classify)
    return d.groupby("Area Category", as_index=False)[area_col].sum()


def _critical_elements(node_df: pd.DataFrame, link_df: pd.DataFrame, units: UnitContext, criteria: ReportCriteria, effective_vel: dict | None = None) -> tuple[pd.DataFrame, pd.DataFrame]:
    node_rows = []
    if node_df is not None and not node_df.empty:
        for _, r in node_df.iterrows():
            node_type = str(r.get("Type", "")).lower()
            # Storage facilities use Calgary storage-classification criteria and are
            # intentionally excluded from generic junction freeboard screening.
            if node_type in {"outfall", "storage"}:
                continue
            ratio = pd.to_numeric(pd.Series([r.get("Depth Ratio")]), errors="coerce").iloc[0]
            free = pd.to_numeric(pd.Series([r.get(f"Freeboard ({units.length})")]), errors="coerce").iloc[0]
            flooding = pd.to_numeric(pd.Series([r.get(f"Peak Flooding ({units.flow})")]), errors="coerce").fillna(0).iloc[0]
            issues = []
            if pd.notna(ratio) and ratio >= criteria.node_depth_ratio:
                issues.append(f"depth ratio ≥ {criteria.node_depth_ratio:.2f}")
            if pd.notna(free) and free <= criteria.minimum_freeboard:
                issues.append(f"freeboard ≤ {criteria.minimum_freeboard:g} {units.length}")
            if flooding > 0:
                issues.append("flooding reported")
            if issues:
                node_rows.append({"Node ID": r.get("Node ID"), "Type": r.get("Type"), "Depth Ratio": ratio,
                                  f"Freeboard ({units.length})": free, f"Peak Flooding ({units.flow})": flooding,
                                  "Review Issue": "; ".join(issues)})
    link_rows = []
    if link_df is not None and not link_df.empty:
        vel_col = f"Peak Velocity ({units.velocity})"
        for _, r in link_df.iterrows():
            if str(r.get("Type", r.get("Model Type", ""))).lower() not in ("conduit", ""):
                continue
            ratio = pd.to_numeric(pd.Series([r.get("Depth Ratio")]), errors="coerce").iloc[0]
            vel = pd.to_numeric(pd.Series([r.get(vel_col)]), errors="coerce").iloc[0]
            source_note = ""
            eff = effective_vel.get(str(r.get("Link ID"))) if effective_vel else None
            if eff is not None:
                if eff.get("Screening Velocity (m/s)") is not None:
                    if "rpt" in str(eff.get("Evidence Source", "")):
                        vel = eff["Screening Velocity (m/s)"]
                        source_note = " [engine .rpt value governs - reconciliation-flagged]"
            issues = []
            advisory = getattr(criteria, "velocity_advisory", 3.0)
            if pd.notna(vel) and vel > criteria.velocity_threshold:
                issues.append(f"CRITICAL screening exceedance: velocity > {criteria.velocity_threshold:g} {units.velocity}{source_note}")
            elif pd.notna(vel) and vel > advisory:
                issues.append(f"Advisory screening exceedance: velocity > {advisory:g} {units.velocity}{source_note}")
            if pd.notna(ratio) and ratio >= criteria.conduit_depth_ratio:
                issues.append(f"depth ratio ≥ {criteria.conduit_depth_ratio:.2f}")
            if issues:
                link_rows.append({"Link ID": r.get("Link ID"), "From Node": r.get("From Node"), "To Node": r.get("To Node"),
                                  f"Peak Flow ({units.flow})": r.get(f"Peak Flow ({units.flow})"), vel_col: vel,
                                  "Depth Ratio": ratio, "Review Issue": "; ".join(issues)})
    return pd.DataFrame(node_rows), pd.DataFrame(link_rows)

def _summary_findings(node_df: pd.DataFrame, link_df: pd.DataFrame, sub_df: pd.DataFrame, metadata: dict[str, Any], units: UnitContext, options: dict[str, str], criteria: ReportCriteria) -> list[str]:
    findings: list[str] = []
    from screening_logic import continuity_disclosure, execution_integrity_assessment
    integrity = execution_integrity_assessment(metadata)
    if not integrity["results_usable"]:
        findings.append("HYDRAULIC RESULTS INVALID: " + integrity["reason"])
    findings.extend(continuity_disclosure(
        metadata, review_pct=criteria.continuity_review,
        warning_pct=criteria.continuity_warning,
        has_pollutants=bool(metadata.get("has_pollutants"))))
    if False:  # legacy block replaced by deterministic continuity_disclosure
        pass
    if not integrity["results_usable"]:
        findings.append("All result-dependent hydraulic screening is Not assessed; raw arrays remain in the audit package only.")
        return findings
    if node_df is not None and not node_df.empty:
        flooded_col = next((c for c in node_df.columns if c.startswith('Peak Flooding (')), None)
        flooded = int((pd.to_numeric(node_df[flooded_col], errors='coerce').fillna(0) > 0).sum()) if flooded_col else 0
        findings.append(f"Flooded nodes identified: {flooded}.")
        junction_nodes = node_df[node_df.get('Type', '').astype(str).str.lower().isin(['junction', 'divider'])].copy()
        if 'Depth Ratio' in junction_nodes and not junction_nodes.empty:
            ratios = pd.to_numeric(junction_nodes['Depth Ratio'], errors='coerce')
            if ratios.notna().any():
                i = ratios.idxmax(); findings.append(f"Maximum junction depth ratio: {ratios.loc[i]:.3f} at {junction_nodes.loc[i, 'Node ID']}.")
                free_col = f"Freeboard ({units.length})"
                if ratios.loc[i] >= criteria.node_depth_ratio and free_col in junction_nodes:
                    findings.append(f"{junction_nodes.loc[i, 'Node ID']} has {_format_value(junction_nodes.loc[i, free_col])} {units.length} of modelled rim clearance and should be reviewed.")
        storage_nodes = node_df[node_df.get('Type', '').astype(str).str.lower().eq('storage')].copy()
        if 'Depth Ratio' in storage_nodes and not storage_nodes.empty:
            ratios = pd.to_numeric(storage_nodes['Depth Ratio'], errors='coerce')
            if ratios.notna().any():
                i = ratios.idxmax(); findings.append(f"Maximum storage depth utilization ratio: {ratios.loc[i]:.3f} at {storage_nodes.loc[i, 'Node ID']}; assess using the applicable storage classification and ponding-depth criterion.")
    conduits = link_df.copy() if link_df is not None else pd.DataFrame()
    if not conduits.empty and 'Type' in conduits:
        conduits = conduits[conduits['Type'].astype(str).str.lower().eq('conduit')]
    if not conduits.empty:
        if "Depth Ratio" in conduits:
            ratios = pd.to_numeric(conduits["Depth Ratio"], errors="coerce")
            if ratios.notna().any():
                i=ratios.idxmax(); findings.append(f"Maximum conduit depth ratio: {ratios.loc[i]:.3f} at {conduits.loc[i, 'Link ID']}.")
        vcol=f"Peak Velocity ({units.velocity})"
        if vcol in conduits:
            vals=pd.to_numeric(conduits[vcol], errors='coerce')
            if vals.notna().any():
                i=vals.idxmax(); findings.append(f"Maximum conduit velocity: {vals.loc[i]:.3f} {units.velocity} at {conduits.loc[i, 'Link ID']}.")
    if sub_df is not None and not sub_df.empty:
        acol=next((c for c in sub_df.columns if c.startswith('Area (')), None)
        if acol: findings.append(f"Total model-database subcatchment area: {pd.to_numeric(sub_df[acol], errors='coerce').sum():.3f} {units.area}.")
    routing = options.get("FLOW_ROUTING", "").upper()
    if routing == "KINWAVE":
        findings.append("⚠️ Kinematic-wave routing does not fully represent backwater, pressurization, reverse flow, or complex surcharge interactions; HGL and surcharge conclusions should be interpreted accordingly.")
    warnings = metadata.get("warnings") or []
    if warnings: findings.append(f"Simulation warnings recorded: {len(warnings)}. Review the attached metadata and model report.")
    return findings


def _event_summary(sub_df: pd.DataFrame, metadata: dict[str, Any], options: dict[str, str], units: UnitContext) -> pd.DataFrame:
    rain_total_col = next((c for c in sub_df.columns if c.startswith("Total Rainfall (")), None) if sub_df is not None else None
    rain_peak_col = next((c for c in sub_df.columns if c.startswith("Peak Rainfall (")), None) if sub_df is not None else None
    total_rain = pd.to_numeric(sub_df[rain_total_col], errors="coerce").max() if rain_total_col and not sub_df.empty else None
    peak_rain = pd.to_numeric(sub_df[rain_peak_col], errors="coerce").max() if rain_peak_col and not sub_df.empty else None
    start = metadata.get("start_time")
    end = metadata.get("end_time")
    duration = ""
    try:
        duration = str(pd.Timestamp(end) - pd.Timestamp(start)) if start and end else ""
    except Exception:
        pass
    rows = [
        {"Parameter": "Design event", "Value": metadata.get("design_storm", "Model design event")},
        {"Parameter": "Simulation start", "Value": start or "Not identified"},
        {"Parameter": "Simulation end", "Value": end or "Not identified"},
        {"Parameter": "Simulation duration", "Value": duration or "Not identified"},
        {"Parameter": "Rainfall / wet-weather timestep", "Value": options.get("WET_STEP", options.get("REPORT_STEP", "Not identified"))},
        {"Parameter": "Reporting timestep", "Value": options.get("REPORT_STEP", "Not identified")},
    ]
    if total_rain is not None and pd.notna(total_rain): rows.append({"Parameter": "Total event precipitation", "Value": f"{_format_value(total_rain)} {'in' if units.system == 'US Customary' else 'mm'}"})
    if peak_rain is not None and pd.notna(peak_rain):
        wet_step = str(options.get("WET_STEP", options.get("REPORT_STEP", "")))
        label = "Maximum rainfall-interval intensity"
        if wet_step in {"0:05", "00:05", "0:05:00", "00:05:00", "5 min", "5 minutes"}:
            label = "Maximum 5-minute rainfall intensity"
        rows.append({"Parameter": label, "Value": f"{_format_value(peak_rain)} {units.rainfall}"})
    return pd.DataFrame(rows)


def _executive_summary(findings: list[str], critical_nodes: pd.DataFrame, critical_links: pd.DataFrame, units: UnitContext, criteria: ReportCriteria) -> list[str]:
    invalid_line = next((x for x in findings if x.startswith("HYDRAULIC RESULTS INVALID:")), None)
    if invalid_line:
        return [
            "The engine process completed, but the hydraulic routing solution failed the deterministic execution-integrity gate.",
            invalid_line,
            "Hydraulic arrays are retained for audit only. Capacity, surcharge, flooding, storage, control, spill-route and depth-velocity conclusions are Not assessed.",
        ]
    lines = ["The simulation completed and the principal model results were screened using the project criteria listed in this report."]
    if critical_nodes.empty:
        lines.append(f"No junctions or dividers met the generic node screening criteria (depth ratio ≥ {criteria.node_depth_ratio:.2f}, rim clearance ≤ {criteria.minimum_freeboard:g} {units.length}, or flooding greater than zero). Storage facilities are assessed separately using their Calgary storage classification.")
    else:
        lines.append(f"{len(critical_nodes)} junction/divider node(s) met the generic node screening criteria; storage facilities are assessed separately using their Calgary storage classification.")
    if critical_links.empty:
        lines.append(f"No conduits met the critical-link screening criteria (depth ratio ≥ {criteria.conduit_depth_ratio:.2f} or velocity ≥ {criteria.velocity_threshold:g} {units.velocity}).")
    else:
        lines.append(f"{len(critical_links)} conduit(s) met the critical-link screening criteria, primarily due to velocity or depth utilization.")
    warning_lines = [x for x in findings if "⚠️" in x]
    if warning_lines:
        lines.append("At least one numerical or modelling limitation requires review before the report is relied upon for design conclusions.")
    return lines


def _provided(value: Any) -> bool:
    return bool(str(value or "").strip()) and str(value).strip().lower() not in {"not provided", "n/a", "none", "unknown"}

def _build_swmr_checklist(metadata: ReportMetadata, criteria: ReportCriteria, *, has_model: bool, has_results: bool, has_storage: bool, has_controls: bool, has_overland: bool, has_outfalls: bool, has_conduits: bool) -> pd.DataFrame:
    drawings = {x.strip().lower() for x in criteria.drawing_inventory}
    reports = list(criteria.applicable_reports)
    rows = []
    def add(i, req, status, evidence, action, category):
        override=(criteria.checklist_overrides or {}).get(i)
        rows.append({"Item":i,"Category":category,"Requirement":req,"Status":override or status,"Report Evidence":evidence,"Outstanding Action":action})
    admin_ok=all(_provided(x) for x in [metadata.project_name, metadata.client, metadata.consultant, metadata.outline_plan_no, metadata.prepared_by])
    add("SWMR-01","Project, developer, consultant, planning and professional information", "Complete" if admin_ok else "Partially complete", "Cover page", "Complete missing administrative and professional fields", "Administration")
    add("SWMR-02","Cover letter, circulation status, unresolved matters and departures", "Partially complete", "Outstanding Information and Actions", "Prepare signed cover letter and identify unresolved matters", "Administration")
    add("SWMR-03","Applicable MDP/SMDP, pond report, prior SWMR and downstream reports", "Complete" if reports else "Missing", "Applicable Reports Register" if reports else "—", "Upload or list applicable drainage documents", "Criteria")
    add("SWMR-04","Study area, legal description, adjacent lands, external drainage and location figure", "Partially complete" if _provided(metadata.legal_description) else "Missing", "Section 2 / model area table", "Add legal description, site location, external drainage and study-area figure", "Site")
    add("SWMR-05","Design objectives and verified project criteria", "Partially complete", "Project Criteria Register", "Confirm current City amendments and project-specific criteria", "Criteria")
    add("SWMR-06","Model methodology, software, routing, infiltration, storm, timesteps and continuity", "Complete" if has_model and has_results else "Missing", "Section 3 and design-event table", "Provide model/results" if not has_results else "Professional confirmation", "Methodology")
    add("SWMR-07","Subcatchment boundaries, areas, imperviousness, widths, slopes and outlets", "Complete" if has_model else "Missing", "Catchment tables", "Confirm against drainage drawings", "Hydrology")
    if has_conduits:
        add("SWMR-08","Minor-system routed flows, cumulative design flows, pipe capacities and spare capacity", "Partially complete" if has_results else "Missing", "Minor-system tables", "Verify release rate and full-flow capacities", "Minor System")
        add("SWMR-09","HGL, surcharge, rim clearance, downstream HWL and backwater assessment", "Partially complete" if has_results else "Missing", "Node HGL table", "Confirm downstream HWL and prepare profiles where required", "Minor System")
    else:
        add("SWMR-08","Minor-system routed flows, cumulative design flows, pipe capacities and spare capacity", "Not applicable", "No conduits identified in uploaded model", "Confirm whether a separate minor-system model is within the report scope", "Minor System")
        add("SWMR-09","HGL, surcharge, rim clearance, downstream HWL and backwater assessment", "Requires professional confirmation", "Storage/outfall HGL table; no conduit network identified", "Confirm whether a separate minor-system model and downstream HWL assessment are required", "Minor System")
    add("SWMR-10","Catchbasin, inlet, ICD and outlet rating information", "Complete" if has_controls else "Not applicable / missing", "Hydraulic controls table" if has_controls else "—", "Confirm inlet types, rating curves and drawing locations", "Minor System")
    add("SWMR-11","Critical overland flows, depths, velocities, spill routes and escape routes", "Partially complete" if has_overland else "Requires drawing review", "Table 9 and deterministic Figure 9-1 depth-velocity screen" if has_overland else "Figure 9-1 criterion included; no overland route identified in model", "Confirm overland route, grading, containment, escape routes and safety on drawings", "Major System")
    add("SWMR-12","Trap-low and storage volume, depth, spill, entrance grade and restrictive covenant information", "Partially complete" if has_storage else "Not applicable", "Storage assessment" if has_storage else "—", "Add spill elevations, building grades and covenant requirements", "Storage")
    add("SWMR-13","Minor and major boundary inflows/outflows and supporting source", "Partially complete" if has_outfalls else "Missing", "Boundary outflow tables" if has_outfalls else "—", "Confirm downstream capacity and external inflows", "Boundary Conditions")
    add("SWMR-14","Private-site permissible discharge and on-site storage requirements", "Missing", "—", "Provide applicable release rates and private-site storage criteria", "Private Sites")
    add("SWMR-15","Water-quality treatment, BMPs, downstream treatment and source controls", "Missing", "—", "Document water-quality strategy and downstream treatment", "Water Quality")
    required_drawings={"site location","study area","catchment plan","model schematic","overland drainage","storm drainage"}
    found=len(required_drawings & drawings)
    if has_model and "model schematic" not in drawings:
        found += 1  # deterministic Figure 3-1 generated from the uploaded INP
    add("SWMR-16","Required figures and drawings", "Complete" if found==len(required_drawings) else ("Partially complete" if found else "Missing"), f"{found}/{len(required_drawings)} core figures/drawings recorded or generated", "Add missing site, catchment, overland and storm-drainage drawings; reconcile Figure 3-1 to issued drawings", "Drawings")
    add("SWMR-17","Model input/output files, formatted listings, model schematic, drawing reconciliation and auditable digital package", "Partially complete" if has_model and has_results else "Missing", "Digital model package, Figure 3-1 and structured appendices" if has_model else "—", "Complete drawing-to-model cross-reference, revision metadata and final authenticated files", "Appendices")
    return pd.DataFrame(rows)

def _readiness_scores(checklist: pd.DataFrame) -> pd.DataFrame:
    weights={"Complete":1.0,"Partially complete":0.5,"Requires professional confirmation":0.5,"Requires drawing review":0.5,"Not applicable":1.0,"Not applicable / missing":0.5,"Missing":0.0}
    cats={"Model data completeness":["Methodology","Hydrology","Appendices"],"Hydraulic-result completeness":["Minor System","Major System","Storage","Boundary Conditions"],"Project information":["Administration","Site"],"Drawing completeness":["Drawings"],"Criteria verification":["Criteria","Private Sites","Water Quality"]}
    rows=[]
    for name,groups in cats.items():
        d=checklist[checklist["Category"].isin(groups)]
        score=100*sum(weights.get(str(x),0.25) for x in d["Status"])/max(len(d),1)
        rows.append({"Readiness Dimension":name,"Score (%)":round(score,1)})
    rows.append({"Readiness Dimension":"SWMR draft readiness","Score (%)":round(sum(r["Score (%)"] for r in rows)/len(rows),1)})
    return pd.DataFrame(rows)

def _add_narrative_block(doc: Document, narrative_sections: Mapping[str, str] | None, key: str) -> bool:
    """Insert an approved narrative block while preserving simple paragraph/list structure."""
    if not narrative_sections:
        return False
    text = str(narrative_sections.get(key, "") or "").strip()
    if not text:
        return False
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        # Section headings are controlled by the deterministic report template.
        if line.startswith("#"):
            continue
        if line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            doc.add_paragraph(re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
        else:
            doc.add_paragraph(line)
    return True



def _scenario_inp_sections(record: Mapping[str, Any]) -> dict[str, list[list[str]]]:
    raw = (record.get("files", {}) or {}).get("inp", b"")
    if isinstance(raw, bytes):
        text = raw.decode("utf-8", errors="ignore")
    else:
        text = str(raw or "")
    sections: dict[str, list[list[str]]] = {}
    current = None
    for original in text.splitlines():
        line = original.strip()
        if line.startswith("[") and line.endswith("]"):
            current = line[1:-1].strip().upper()
            sections.setdefault(current, [])
            continue
        if not current or not line or line.startswith(";"):
            continue
        data = line.split(";", 1)[0].strip()
        if data:
            sections[current].append(data.split())
    return sections

def _series_peak(values: Mapping[str, Any], key: str, absolute: bool = False) -> float:
    seq = values.get(key, []) or []
    nums=[]
    for v in seq:
        try:
            x=float(v); nums.append(abs(x) if absolute else x)
        except (TypeError, ValueError):
            pass
    return max(nums, default=0.0)

def _safe_float(value: Any) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _peak_time(values: Mapping[str, Any], key: str, times: list[Any]) -> str:
    seq = values.get(key, []) or []
    if not seq:
        return "Not available"
    numeric=[]
    for i, value in enumerate(seq):
        try:
            numeric.append((float(value), i))
        except (TypeError, ValueError):
            continue
    if not numeric:
        return "Not available"
    _, idx=max(numeric, key=lambda x: x[0])
    if idx < len(times):
        return str(times[idx])
    return str(idx)


def _event_short_label(row: Mapping[str, Any]) -> str:
    role=str(row.get("Model Role", "Scenario") or "Scenario")
    storm=str(row.get("Storm", "Model event") or "Model event")
    low=storm.lower().replace("_", " ")
    match=re.search(r"(?:^|\D)(\d+)\s*(?:y|yr|year)", low)
    if not match:
        match=re.search(r"(\d+)y", low.replace(" ", ""))
    event=f"{match.group(1)}-Year" if match else storm
    prefix="Base" if role.lower()=="base" else "Scenario"
    return f"{prefix} – {event}"


def _scenario_detail_tables(record: Mapping[str, Any], units: UnitContext) -> dict[str, pd.DataFrame]:
    results = record.get("results", {}) or {}
    definition = record.get("definition", {}) or {}
    summary = record.get("summary", {}) or {}
    sections = _scenario_inp_sections(record)
    storage_rows_inp = {r[0]: r for r in sections.get("STORAGE", []) if r}
    storage_ids = set(storage_rows_inp)
    outfall_ids = {r[0] for r in sections.get("OUTFALLS", []) if r}
    conduit_ids = {r[0] for r in sections.get("CONDUITS", []) if r}
    outlet_ids = {r[0] for r in sections.get("OUTLETS", []) if r}
    node_ts = results.get("node_ts", {}) or {}
    link_ts = results.get("link_ts", {}) or {}
    sub_ts = results.get("sub_ts", {}) or {}
    times = list(results.get("times", []) or [])

    overview = pd.DataFrame([{
        "Scenario": definition.get("scenario_name", summary.get("Scenario Name", "Scenario")),
        "Source Model": summary.get("Source Model", (record.get("manifest", {}) or {}).get("base_model_name", "Not identified")),
        "Rainfall Event": (definition.get("storm", {}) or {}).get("name", summary.get("Storm", "Model rainfall")),
        "Storm Status": summary.get("Storm Status", (definition.get("storm", {}) or {}).get("source_status", "Not verified")),
        "Runoff Error (%)": summary.get("Runoff Error (%)"),
        "Flow Error (%)": summary.get("Flow Error (%)"),
        f"Peak Subcatchment Runoff ({units.flow})": summary.get("Peak Subcatchment Runoff"),
        f"Peak Link/Control Flow ({units.flow})": summary.get("Peak Link Flow"),
        f"Maximum Node Inflow ({units.flow})": summary.get("Maximum Node Inflow"),
        f"Maximum Node Flooding ({units.flow})": summary.get("Maximum Node Flooding"),
        "Maximum Conduit Velocity": (summary.get("Maximum Link Velocity") if conduit_ids else "Not applicable — no conduits in model"),
        "Maximum Conduit Depth Ratio": (summary.get("Maximum Modelled Depth Ratio") if conduit_ids else "Not applicable — no conduits in model"),
    }])

    sub_rows=[]
    for sid, vals in sub_ts.items():
        sub_rows.append({"Subcatchment": sid, f"Peak Runoff ({units.flow})": _series_peak(vals,"runoff"), f"Peak Rainfall ({units.rainfall})": _series_peak(vals,"rainfall")})

    storage_rows=[]
    for nid in sorted(storage_ids):
        vals=node_ts.get(nid,{})
        inp=storage_rows_inp.get(nid, [])
        invert=_safe_float(inp[1]) if len(inp)>1 else None
        max_depth=_safe_float(inp[2]) if len(inp)>2 else None
        peak_depth=_series_peak(vals,"depth")
        peak_head=_series_peak(vals,"head")
        if not peak_head and invert is not None:
            peak_head=invert+peak_depth
        peak_volume=_series_peak(vals,"volume")
        remaining=(max_depth-peak_depth) if max_depth is not None else None
        utilization=(100.0*peak_depth/max_depth) if max_depth and max_depth>0 else None
        storage_rows.append({
            "Storage ID":nid,
            f"Peak Depth ({units.length})":peak_depth,
            f"Maximum HGL ({units.length})":peak_head,
            f"Maximum Stored Volume ({units.storage})":peak_volume,
            f"Peak Inflow ({units.flow})":_series_peak(vals,"inflow"),
            f"Peak Outflow ({units.flow})":_series_peak(vals,"outflow"),
            f"Peak Flooding ({units.flow})":_series_peak(vals,"flooding"),
            "Time of Peak Storage":_peak_time(vals,"volume",times),
            "Depth Utilization (%)":utilization,
            f"Modelled Depth Margin ({units.length})":remaining,  # model quantity, NOT regulatory freeboard
        })

    outfall_rows=[]
    for nid in sorted(outfall_ids):
        vals=node_ts.get(nid,{})
        outfall_rows.append({"Outfall ID":nid, f"Peak Depth ({units.length})":_series_peak(vals,"depth"), f"Peak Inflow ({units.flow})":_series_peak(vals,"inflow"), f"Peak Flooding ({units.flow})":_series_peak(vals,"flooding")})
    control_rows=[]
    for lid in sorted(outlet_ids):
        vals=link_ts.get(lid,{})
        control_rows.append({"Control ID":lid, f"Peak Flow ({units.flow})":_series_peak(vals,"flow",True), f"Peak Depth ({units.length})":_series_peak(vals,"depth"), "Time of Peak Flow":_peak_time(vals,"flow",times)})
    conduit_rows=[]
    for lid in sorted(conduit_ids):
        vals=link_ts.get(lid,{})
        diameter=float(vals.get("diameter",0) or 0)
        depth=_series_peak(vals,"depth")
        conduit_rows.append({"Conduit ID":lid, f"Peak Flow ({units.flow})":_series_peak(vals,"flow",True), f"Peak Velocity ({units.velocity})":_series_peak(vals,"velocity",True), "Modelled Depth Ratio": depth/diameter if diameter>0 else None})
    return {"overview":overview,"subcatchments":pd.DataFrame(sub_rows),"storage":pd.DataFrame(storage_rows),"outfalls":pd.DataFrame(outfall_rows),"controls":pd.DataFrame(control_rows),"conduits":pd.DataFrame(conduit_rows)}

def _compact_scenario_comparison(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    out=df.copy()
    out.insert(0, "Display Scenario", [_event_short_label(row) for _, row in out.iterrows()])
    wanted=["Display Scenario","Source Model","Storm","Storm Status","Simulation Status","Runoff Error (%)","Flow Error (%)","Peak Subcatchment Runoff","Maximum Storage Depth","Maximum Storage Volume","Peak Link Flow","Maximum Node Inflow","Maximum Node Flooding","Hydraulic Difference"]
    return out[[c for c in wanted if c in out.columns]].copy()

_LISTING_MAX_LINES = 8000
_LISTING_CHUNK = 50


def _add_fixed_width_listing(doc, heading: str, text: str, level: int = 2,
                             max_lines: int = _LISTING_MAX_LINES) -> None:
    """Append a fixed-width (Courier) model-file listing under a heading.

    Long files are truncated in the MIDDLE (head + tail preserved) with an
    explicit note directing readers to the untruncated copy in the audit
    package. Lines are batched ~50 per paragraph via run breaks to keep the
    document responsive in Word.
    """
    from docx.shared import Pt

    doc.add_heading(heading, level=level)
    lines = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    total = len(lines)
    if total > max_lines:
        head_n = int(max_lines * 0.7)
        tail_n = max_lines - head_n
        omitted = total - head_n - tail_n
        lines = (lines[:head_n]
                 + [f"... [{omitted} lines omitted for document size - complete file "
                    f"is included untruncated in the audit package under model/] ..."]
                 + lines[-tail_n:])
        doc.add_paragraph(
            f"Listing truncated for document size ({total} source lines; {omitted} omitted "
            "mid-file). The complete, untruncated file is archived in the audit package "
            "(model/ directory) and is the authoritative record.")
    for start in range(0, len(lines), _LISTING_CHUNK):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run()
        run.font.name = "Courier New"
        run.font.size = Pt(6.5)
        chunk = lines[start:start + _LISTING_CHUNK]
        for j, line in enumerate(chunk):
            if j:
                run.add_break()
            run.add_text(line if line.strip() else " ")


def _embed_attached_figures(doc, figures: list[Mapping[str, Any]] | None) -> None:
    """Insert session-attached figures into the built document.

    Each figure dict: {"figure_id", "path", "caption", "section", "source"}.
    Figures are inserted at the END of the first level-1 section whose
    heading contains the requested section keyword (case-insensitive), i.e.
    immediately before the next Heading-1 paragraph; unmatched sections fall
    back to the end of the document. Client-supplied figures are labelled as
    such — they are illustrative material attached to the audited report,
    not server-verified outputs.
    """
    if not figures:
        return
    from docx.shared import Inches

    headings = [(i, p) for i, p in enumerate(doc.paragraphs)
                if p.style.name.startswith("Heading 1")]

    def anchor_for(section_kw: str):
        kw = (section_kw or "results").strip().lower()
        for pos, (idx, para) in enumerate(headings):
            if kw in para.text.lower():
                if pos + 1 < len(headings):
                    return headings[pos + 1][1]  # insert before next H1
                return None  # matched last section -> append at end
        return None

    for n, fig in enumerate(figures, start=1):
        path = str(fig.get("path", ""))
        if not path or not Path(path).exists():
            continue
        caption = str(fig.get("caption") or fig.get("figure_id") or f"Attached figure {n}")
        source = str(fig.get("source") or "session-attached")
        label = f"Figure A{n} - {caption} ({source}; illustrative, not a server-verified output)"
        anchor = anchor_for(str(fig.get("section", "results")))
        try:
            if anchor is not None:
                pic_par = anchor.insert_paragraph_before()
                pic_par.add_run().add_picture(path, width=Inches(6.0))
                cap_par = anchor.insert_paragraph_before(label)
                cap_par.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else cap_par.style
            else:
                doc.add_paragraph().add_run().add_picture(path, width=Inches(6.0))
                doc.add_paragraph(label)
        except Exception:
            # A corrupt image must never abort report generation.
            (anchor.insert_paragraph_before if anchor is not None else doc.add_paragraph)(
                f"[Attached figure '{caption}' could not be embedded — file unreadable.]")


def generate_report_package(
    *,
    metadata: ReportMetadata,
    inp_sections: dict[str, list[list[str]]],
    node_summary: pd.DataFrame,
    link_summary: pd.DataFrame,
    sub_summary: pd.DataFrame,
    simulation_metadata: dict[str, Any],
    result_db_bytes: bytes | None = None,
    criteria: ReportCriteria | None = None,
    narrative_sections: Mapping[str, str] | None = None,
    scenario_comparison: pd.DataFrame | None = None,
    scenario_analysis: str | None = None,
    scenario_records: list[Mapping[str, Any]] | None = None,
    scenario_reporting_mode: str = "Base report with scenario comparison",
    preliminary_review_artifacts: Mapping[str, Any] | None = None,
    attached_figures: list[Mapping[str, Any]] | None = None,
    model_listings: Mapping[str, str] | None = None,
    reconciliation: Mapping[str, Any] | None = None,
    model_identity: Mapping[str, Any] | None = None,
) -> dict[str, bytes | str]:
    """Generate editable Word report and ZIP package entirely in memory."""
    criteria = criteria or ReportCriteria()
    simulation_metadata = dict(simulation_metadata or {})
    from screening_logic import execution_integrity_assessment
    integrity = execution_integrity_assessment(simulation_metadata)
    simulation_metadata.update({
        "execution_integrity_status": integrity["status"],
        "results_usable": integrity["results_usable"],
        "execution_integrity_reason": integrity["reason"],
    })
    options = _inp_options(inp_sections)
    units = _unit_context(options.get("FLOW_UNITS", ""))
    sub_table = _subcatchment_model_table(inp_sections, sub_summary)
    link_table = _link_model_table(inp_sections, link_summary)
    hgl_table = _node_hgl_table(node_summary)

    # The simulation arrays are retained in the model's native unit system.
    # Rename report headers to match that system; do not convert values.
    sub_table = _rename_native_columns(sub_table, units)
    link_table = _rename_native_columns(link_table, units)
    hgl_table = _rename_native_columns(hgl_table, units)
    sub_table = sub_table.loc[:, ~sub_table.columns.duplicated()].copy()
    link_table = link_table.loc[:, ~link_table.columns.duplicated()].copy()
    hgl_table = hgl_table.loc[:, ~hgl_table.columns.duplicated()].copy()
    node_report = _rename_native_columns(node_summary, units)
    link_report = _rename_native_columns(link_summary, units)
    sub_report = _rename_native_columns(sub_summary, units)

    report_step_s = _parse_time_seconds(options.get("REPORT_STEP"))
    sub_report = _enhance_subcatchment_results(sub_report, units, report_step_s)
    sub_table = _enhance_subcatchment_results(sub_table, units, report_step_s)
    node_report = _populate_outfall_flows(node_report, link_report, units)
    db_tables = _read_sqlite_tables(result_db_bytes, ["node_timeseries", "link_timeseries"])
    control_table = _control_table(inp_sections, link_table, units)
    control_recon = _control_reconciliation(control_table, link_table, units)
    storage_table = _storage_table(inp_sections, node_report, db_tables, units)
    if not storage_table.empty and not control_table.empty:
        peak_col = f"Peak Flow ({units.flow})"
        for idx in storage_table.index:
            sid = storage_table.at[idx, "Storage ID"]
            outgoing_peak = pd.to_numeric(control_table.loc[control_table["From Node"].astype(str) == str(sid), peak_col], errors="coerce").fillna(0).sum() if peak_col in control_table else 0.0
            depth_col = next((c for c in storage_table.columns if c.startswith("Time-Series Peak Depth (")), None)
            peak_depth = storage_table.at[idx, depth_col] if depth_col else storage_table.at[idx, f"Reported Peak Depth ({units.length})"]
            if outgoing_peak > 0 and (pd.isna(peak_depth) or float(peak_depth) <= 0):
                storage_table.at[idx, "Status"] = "Review: control flow occurs but storage depth is zero"
    area_table = _area_classification(sub_table, units, criteria.area_classification)
    from screening_logic import effective_velocity_table, missing_information_register
    simulation_metadata["has_pollutants"] = bool(inp_sections.get("POLLUTANTS"))
    recon_links_df = None
    if reconciliation and reconciliation.get("links") is not None:
        recon_links_df = pd.DataFrame(reconciliation["links"]) if not isinstance(reconciliation["links"], pd.DataFrame) else reconciliation["links"]
    effective_table = effective_velocity_table(
        link_report, recon_links_df,
        advisory=criteria.velocity_advisory, critical=criteria.velocity_threshold)
    effective_map = {str(r["Link ID"]): r for r in effective_table.to_dict("records")} if not effective_table.empty else {}
    critical_nodes, critical_links = _critical_elements(_rename_native_columns(_node_hgl_table(node_report), units), link_report, units, criteria, effective_vel=effective_map)

    calgary = CalgaryCriteria(
        minor_release_rate_lps_ha=criteria.minor_release_rate_lps_ha,
        trap_low_max_depth_m=criteria.trap_low_max_depth_m,
        entrance_grade_margin_m=criteria.entrance_grade_margin_m,
        pipe_critical_velocity_mps=criteria.velocity_threshold if units.system == "SI" else 4.0,
        conduit_capacity_review_ratio=criteria.conduit_depth_ratio,
        conduit_capacity_warning_ratio=criteria.conduit_capacity_warning_ratio,
        continuity_review_pct=criteria.continuity_review,
        continuity_warning_pct=criteria.continuity_warning,
        special_link_limits=criteria.special_link_limits or {},
        storage_classification=criteria.storage_classification or {},
        outfall_classification=criteria.outfall_classification or {},
    )
    rain_gages = [r[1] for r in inp_sections.get("SUBCATCHMENTS", []) if len(r) > 1]
    inferred_event, inferred_event_note = infer_design_event(rain_gages, metadata.design_storm)
    if metadata.design_storm.strip().lower() in {"model design event", "not provided", ""}:
        metadata.design_storm = inferred_event
    criteria_table = criteria_register(calgary)
    # Calgary-specific calculations use SI design rules. In US models, tables remain available but are flagged for review.
    minor_capacity = build_minor_system_capacity_table(link_table, units.system, units.flow, units.length, calgary)
    if not integrity["results_usable"]:
        invalid_label = "Not assessed - hydraulic routing solution invalid"
        for df in (node_report, link_report, link_table, control_table, storage_table):
            if df is not None and not df.empty:
                df["Status"] = invalid_label
        if control_recon is not None and not control_recon.empty:
            control_recon["Check"] = "Indeterminate - hydraulic routing solution invalid"
        if minor_capacity is not None and not minor_capacity.empty:
            if "Status" in minor_capacity:
                minor_capacity["Status"] = invalid_label
            if "Assessment Basis" in minor_capacity:
                minor_capacity["Assessment Basis"] = invalid_label
    # Checklist applicability is based on actual model element types, not on
    # whether the generic link table happens to contain outlet controls.
    _conduit_rows = inp_sections.get("CONDUITS", []) or []
    _xsection_rows = {str(r[0]): str(r[1]).upper() for r in (inp_sections.get("XSECTIONS", []) or []) if len(r) > 1}
    _open_shapes = {"TRAPEZOIDAL", "RECT_OPEN", "TRIANGULAR", "IRREGULAR", "STREET"}
    _has_overland = any(str(r[0]) in _xsection_rows and _xsection_rows[str(r[0])] in _open_shapes for r in _conduit_rows if r)
    checklist_table = _build_swmr_checklist(
        metadata,
        criteria,
        has_model=bool(inp_sections),
        has_results=not node_report.empty and integrity["results_usable"],
        has_storage=not storage_table.empty,
        has_controls=not control_table.empty,
        has_overland=_has_overland,
        has_outfalls=bool((node_report.get("Type", pd.Series(dtype=str)).astype(str).str.lower()=="outfall").any()),
        has_conduits=bool(_conduit_rows),
    )
    if model_listings and (model_listings.get("inp_text") or model_listings.get("rpt_text")):
        _m17 = checklist_table["Requirement"].astype(str).str.contains("input/output files", case=False)
        checklist_table.loc[_m17, "Status"] = "Included (Appendix D listings + digital package)"
        checklist_table.loc[_m17, "Outstanding Action"] = "Verify appendix listings match the issued digital model package"
    readiness_table = _readiness_scores(checklist_table)
    schematic_png, schematic_manifest = _model_schematic_figure(inp_sections)

    doc = Document()
    sec = doc.sections[0]
    _footer_p = sec.footer.paragraphs[0]
    _footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run = _footer_p.add_run(f"{metadata.project_name} — Preliminary Draft — Page ")
    _run.font.size = Pt(8)
    _fld = OxmlElement("w:fldSimple"); _fld.set(qn("w:instr"), "PAGE")
    _footer_p._p.append(_fld)
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for sname in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        styles[sname].font.name = "Arial"
        styles[sname].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("STORMWATER MANAGEMENT REPORT")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(31, 78, 121)
    p = doc.add_paragraph(metadata.project_name)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True; p.runs[0].font.size = Pt(16)

    cover = doc.add_table(rows=0, cols=2); cover.style = "Table Grid"
    for label, value in [
        ("Subdivision (SB) #", metadata.subdivision_no), ("Outline Plan #", metadata.outline_plan_no),
        ("Development Permit #", metadata.development_permit_no), ("Prepared for", metadata.client),
        ("Prepared by", metadata.consultant), ("Consultant file number", metadata.consultant_file_no),
        ("Responsible engineer", metadata.prepared_by), ("Checked by", metadata.checked_by),
        ("Contact name", metadata.contact_name), ("Contact email", metadata.contact_email),
        ("Legal description", metadata.legal_description), ("Circulation status", metadata.submission_status),
        ("Construction drawing no.", metadata.construction_drawing_no), ("Development Agreement no.", metadata.development_agreement_no),
        ("Report date", metadata.report_date), ("Reporting profile", metadata.municipality),
    ]:
        cells = cover.add_row().cells
        _set_cell_text(cells[0], label, bold=True, size=9)
        _set_cell_text(cells[1], value, size=9)

    doc.add_page_break()
    doc.add_heading("1.0 INTRODUCTION", level=1)
    if not _add_narrative_block(doc, narrative_sections, "introduction"):
        doc.add_paragraph(
            f"This report summarizes the hydrologic and hydraulic analysis for {metadata.project_name}. "
            "The document was generated from the deterministic EPA SWMM/OpenSWMM simulation results in the SWMM6 GIS Tool. "
            "Project-specific planning, survey, geotechnical, environmental, drawing, and approval information must be verified by the responsible professional engineer."
        )

    doc.add_paragraph("Draft status: suitable for consultant/client discussion and structured engineering review. Human-in-the-loop completion, technical verification, drawing coordination, and professional authentication are required before municipal submission.")
    if model_identity:
        doc.add_heading("1.0a Model Identity and Execution Provenance", level=2)
        identity_rows = [
            {"Item": "Model file", "Value": str(model_identity.get("inp_name", "—"))},
            {"Item": "Model SHA-256", "Value": str(model_identity.get("sha256", "—"))},
            {"Item": "Session / run ID", "Value": f"{model_identity.get('session_id', '—')} / {model_identity.get('run_id', '—')}"},
            {"Item": "Engine", "Value": str(model_identity.get("engine", "EPA SWMM / OpenSWMM (crash-isolated worker)"))},
            {"Item": "Execution status", "Value": str(model_identity.get("status", "Completed"))},
            {"Item": "Generated", "Value": str(model_identity.get("generated", metadata.report_date))},
        ]
        if model_identity.get("legacy_defaults_normalized"):
            identity_rows.extend([
                {"Item": "Execution derivative", "Value": str(model_identity.get("execution_inp_name", "—"))},
                {"Item": "Execution SHA-256", "Value": str(model_identity.get("execution_sha256", "—"))},
                {"Item": "Legacy option handling", "Value": "Zero/default sentinels normalized in an immutable derivative; original upload preserved"},
            ])
        _id_tbl = pd.DataFrame(identity_rows)
        _add_df_table(doc, "Table 0A - Model Identity", _id_tbl, font_size=8.0)
        _rev_tbl = pd.DataFrame([{"Rev": "P0", "Date": metadata.report_date,
                                  "Description": "Preliminary server-generated draft for engineering review",
                                  "Status": metadata.submission_status}])
        _add_df_table(doc, "Table 0B - Revision History", _rev_tbl, font_size=8.0)
    doc.add_heading("1.1 Calgary SWMR Checklist Summary", level=2)
    doc.add_paragraph("The completeness register tracks the principal Calgary SWMR checklist subjects. It is a draft-readiness tool, not a municipal compliance score.")
    _add_df_table(doc, "Table 1A - Calgary SWMR Completeness Register", checklist_table, landscape=True, font_size=7.0)
    _add_df_table(doc, "Table 1B - Draft Readiness Summary", readiness_table, font_size=8.0)
    if criteria.applicable_reports:
        doc.add_heading("1.2 Applicable Reports Register", level=2)
        for item in criteria.applicable_reports: doc.add_paragraph(item, style="List Bullet")
    _add_narrative_block(doc, narrative_sections, "applicable_criteria")

    doc.add_heading("2.0 SITE DESCRIPTION AND DESIGN CRITERIA", level=1)
    _add_narrative_block(doc, narrative_sections, "site_description")
    native_area_col = f"Area ({units.area})"
    total_area = float(sub_report[native_area_col].sum()) if not sub_report.empty and native_area_col in sub_report else 0.0
    weighted_imp = 0.0
    if not sub_report.empty and total_area > 0 and "% Impervious" in sub_report:
        weighted_imp = float((sub_report[native_area_col] * sub_report["% Impervious"]).sum() / total_area)
    doc.add_paragraph(
        f"The model database contains {len(sub_summary) if sub_summary is not None else 0} subcatchments, "
        f"{len(node_summary) if node_summary is not None else 0} nodes, and {len(link_summary) if link_summary is not None else 0} links. "
        f"The combined model-database subcatchment area is {total_area:.3f} {units.area} and the area-weighted imperviousness is approximately {weighted_imp:.1f}%. "
        "Pre-development/comparison and external catchments are listed separately below where identifiable from their names."
    )
    _add_df_table(doc, "Table 2A - Model Area Classification", area_table, font_size=8.2)
    doc.add_heading("2.1 Design Objectives", level=2)
    custom_objectives = str((narrative_sections or {}).get("design_objectives", "") or "").strip()
    if custom_objectives:
        for line in [l.strip().lstrip("-*\u2022 ").strip() for l in custom_objectives.splitlines() if l.strip()]:
            doc.add_paragraph(line, style="List Bullet")
    else:
        for text in [
            "Confirm minor-system flows remain within selected hydraulic criteria.",
            "Assess overland conveyance depth and velocity against the applicable municipal criteria.",
            "Identify surcharge, flooding, instability, and continuity concerns.",
            "Confirm boundary outflows and runoff volumes for the selected design event.",
        ]:
            doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("3.0 ANALYSIS METHODOLOGY AND DATA", level=1)
    _add_narrative_block(doc, narrative_sections, "methodology")
    doc.add_heading("3.1 Design Storm", level=2)
    simulation_metadata = dict(simulation_metadata or {})
    simulation_metadata["design_storm"] = metadata.design_storm
    event_table = _event_summary(sub_report, simulation_metadata, options, units)
    _add_df_table(doc, "Table 3A - Design Event Summary", event_table, font_size=8.2)
    doc.add_paragraph(inferred_event_note)
    if criteria.calgary_enabled:
        _add_df_table(doc, "Table 3B - Calgary Project Criteria Register", criteria_table, landscape=True, font_size=7.5)
    doc.add_heading("3.2 Computer Model", level=2)
    doc.add_paragraph(
        f"The analysis was completed using EPA SWMM/OpenSWMM. The report retains the model unit system ({units.system}; flow units {units.flow}). "
        f"routing model: {options.get('FLOW_ROUTING', 'not identified')}; infiltration model: {options.get('INFILTRATION', 'not identified')}; "
        f"reporting step: {options.get('REPORT_STEP', 'not identified')}; routing step: {options.get('ROUTING_STEP', 'not identified')}."
    )
    if options.get("FLOW_ROUTING", "").upper() == "KINWAVE":
        doc.add_paragraph("Modelling limitation: kinematic-wave routing does not fully represent backwater, pressurization, reverse flow, or complex surcharge interactions. Dynamic-wave routing should be considered where these effects are material.")
    doc.add_picture(io.BytesIO(schematic_png), width=Inches(7.15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    schematic_caption = doc.add_paragraph(
        "Figure 3-1 - Automated SWMM Model Schematic. Solid arrows show hydraulic-link "
        "direction and dashed arrows show subcatchment runoff routing. This topology figure "
        "is generated from the uploaded INP and must be reconciled to the issued drainage drawings."
    )
    schematic_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in schematic_caption.runs:
        run.italic = True
        run.font.size = Pt(8)
    doc.add_heading("3.3 Major-Minor System", level=2)
    doc.add_paragraph("Open channels, swales, gutters, culverts, and closed conduits represented in the model were reviewed. Final major/minor classification must be confirmed against the approved drainage concept and drawings.")
    if scenario_comparison is not None and not scenario_comparison.empty:
        doc.add_heading("3.4 Preliminary Design Scenarios", level=2)
        doc.add_paragraph(
            "The scenario register summarizes preliminary model alternatives generated and simulated through the Rev22 workflow. "
            "Scenario storm sources, parameter changes, and review status must be confirmed by the responsible engineer before design use."
        )
        doc.add_paragraph(f"Reporting mode: {scenario_reporting_mode}.")
        compact_scenario = _compact_scenario_comparison(scenario_comparison)
        _add_df_table(doc, "Table 3C - Preliminary Scenario Comparison", compact_scenario, landscape=True, font_size=7.2)
        if scenario_analysis:
            doc.add_heading("3.4.1 Scenario Comparison Analysis", level=3)
            for paragraph in str(scenario_analysis).split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
        if scenario_records and scenario_reporting_mode != "Base report with scenario comparison":
            doc.add_heading("3.4.2 Event-Specific Deterministic Results", level=3)
            doc.add_paragraph("Each selected model-event combination is presented as a separate deterministic dataset. Values shown as not applicable indicate that the corresponding model element type is absent.")
            table_no = 1
            for record in scenario_records:
                definition = record.get("definition", {}) or {}
                summary = record.get("summary", {}) or {}
                label = _event_short_label(summary)
                full_label = definition.get("scenario_name", f"Scenario {table_no}")
                doc.add_heading(str(label), level=4)
                doc.add_paragraph(f"Model-event dataset: {full_label}. Source model: {summary.get('Source Model', 'Not identified')}. Rainfall event: {summary.get('Storm', 'Not identified')}.")
                detail = _scenario_detail_tables(record, units)
                for key, title in [("overview","Event Summary"),("storage","Storage Results"),("controls","Outlet and Control Results"),("outfalls","Outfall Results"),("conduits","Conduit Results"),("subcatchments","Subcatchment Results")]:
                    frame=detail[key]
                    if frame is not None and not frame.empty:
                        _add_df_table(doc, f"Table 3D-{table_no} - {label}: {title}", frame, landscape=len(frame.columns)>6, font_size=7.4)
                        table_no += 1
        doc.add_heading("3.5 Catchment Areas", level=2)
    else:
        doc.add_heading("3.4 Catchment Areas", level=2)
    _add_narrative_block(doc, narrative_sections, "hydrology")
    catchment_main_cols = [
        "Sub ID", "Rain Gage", "Outlet", f"Area ({units.area})", "Impervious (%)",
        f"Width ({units.length})", "Slope (%)", "n Imperv.", "n Perv.",
        f"Peak Runoff ({units.flow})", f"Peak Rainfall ({units.rainfall})", "Peak-Flow Runoff Coefficient",
    ]
    _add_df_table(doc, "Table 4 - Catchment Parameters and Runoff Results", sub_table[[c for c in catchment_main_cols if c in sub_table.columns]], landscape=True)

    doc.add_heading("4.0 RESULTS", level=1)
    doc.add_heading("4.1 Major-System and Spill-Route Assessment", level=2)
    _add_narrative_block(doc, narrative_sections, "major_system")
    overland = link_table.copy()
    if criteria.major_link_ids:
        selected = {str(x).strip() for x in criteria.major_link_ids if str(x).strip()}
        overland = overland[overland["Link ID"].astype(str).isin(selected)].copy()
    else:
        # Outlet/orifice/weir controls are not automatically classified as
        # major-system overland routes. Only open conduit/street shapes are
        # included unless the engineer explicitly supplies major_link_ids.
        model_type = overland.get("Model Type", overland.get("Type", pd.Series("", index=overland.index))).astype(str).str.lower()
        shape = overland.get("Shape", pd.Series("", index=overland.index)).fillna("").astype(str).str.upper()
        overland = overland[model_type.eq("conduit") & shape.isin(["TRAPEZOIDAL", "RECT_OPEN", "TRIANGULAR", "IRREGULAR", "STREET"])].copy()
    overland_cols = ["Link ID", "From Node", "To Node", "Shape", f"Peak Flow ({units.flow})", f"Peak Depth ({units.length})", f"Peak Velocity ({units.velocity})", "Depth Ratio", "Status"]
    _add_df_table(doc, "Table 9 - Overland Flow Assessment", overland[[c for c in overland_cols if c in overland.columns]], landscape=True)
    overland_compliance = build_overland_compliance_table(overland, calgary, units.flow, units.length, units.velocity) if units.system == "SI" else pd.DataFrame()
    if not integrity["results_usable"] and not overland_compliance.empty:
        for col in ("Depth-Velocity Status", "Special Limit Status"):
            if col in overland_compliance:
                overland_compliance[col] = "Not assessed - hydraulic routing solution invalid"
        if "Spill Active" in overland_compliance:
            overland_compliance["Spill Active"] = "Not assessed"
    depth_velocity_png = _depth_velocity_figure(
        overland_compliance, calgary.depth_velocity_curve,
        results_usable=integrity["results_usable"], flow_unit=units.flow)
    doc.add_picture(io.BytesIO(depth_velocity_png), width=Inches(6.65))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph(
        "Figure 9-1 - Alberta/Calgary Depth–Velocity Criteria for Overland Flow. "
        "Model points are derived from Table 9; marker colour represents peak flow. "
        "Straight-line interpolation is used between the tabulated points. The 2011 baseline "
        "envelope must be checked against current and project-specific requirements."
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.italic = True
        run.font.size = Pt(8)
    curve_table = pd.DataFrame(calgary.depth_velocity_curve,
                               columns=["Water Velocity (m/s)", "Permissible Depth (m)"])
    _add_df_table(doc, "Figure 9-1A - Permissible Depth and Velocity of Overland Flow",
                  curve_table, landscape=False, font_size=8.0)
    if not overland_compliance.empty:
        _add_df_table(doc, "Table 9A - Calgary Major-System Depth-Velocity Screening", overland_compliance, landscape=True, font_size=7.2)

    outfalls = node_report[node_report["Type"].astype(str).str.lower() == "outfall"].copy() if not node_report.empty and "Type" in node_report else pd.DataFrame()
    if not outfalls.empty:
        outfalls["Status"] = "Boundary condition"
        outfalls["Depth Ratio"] = "N/A"
    outfall_cols = ["Node ID", "Type", f"Invert ({units.length})", f"Peak Depth ({units.length})", f"Peak Inflow ({units.flow})", f"Peak Flooding ({units.flow})", "Status"]
    _add_df_table(doc, "Table 10 - Major System Boundary Conditions - Outflows", outfalls[[c for c in outfall_cols if c in outfalls.columns]])

    doc.add_heading("4.2 Minor-System Assessment", level=2)
    _add_narrative_block(doc, narrative_sections, "minor_system")
    if f"Diameter ({units.length})" in link_table.columns:
        link_table = link_table.rename(columns={f"Diameter ({units.length})": f"Diameter / Geom1 ({units.length})"})
    minor_cols = ["Link ID", "From Node", "To Node", "Shape", f"Length ({units.length})", f"Diameter / Geom1 ({units.length})", "Manning n", f"Peak Flow ({units.flow})", f"Peak Depth ({units.length})", "Depth Ratio", f"Peak Velocity ({units.velocity})", "Status"]
    conduit_table = link_table[link_table.get("Model Type", link_table.get("Type", "")).astype(str).str.lower().eq("conduit")].copy()
    _add_df_table(doc, "Table 12A - Conduit, Culvert, Swale and Gutter Analysis", conduit_table[[c for c in minor_cols if c in conduit_table.columns]], landscape=True)
    if not minor_capacity.empty:
        _add_df_table(doc, "Table 12A-1 - Calgary Minor-System Capacity Screening", minor_capacity, landscape=True, font_size=7.2)
    doc.add_heading("4.3 Hydraulic Controls", level=2)
    _add_narrative_block(doc, narrative_sections, "hydraulic_controls")
    if not criteria.suppress_empty_sections or not control_table.empty:
        _add_df_table(doc, "Table 12B - Hydraulic Controls", control_table, landscape=True, font_size=7.2)
    if not criteria.suppress_empty_sections or not control_recon.empty:
        _add_df_table(doc, "Table 12C - Hydraulic Control Flow Reconciliation", control_recon, landscape=True, font_size=7.2)
    doc.add_heading("4.4 Storage and Trap-Low Assessment", level=2)
    _add_narrative_block(doc, narrative_sections, "storage")
    storage_calgary = apply_storage_classification(storage_table, calgary, units.length) if not storage_table.empty else pd.DataFrame()
    if not integrity["results_usable"] and not storage_calgary.empty:
        if "Calgary Status" in storage_calgary:
            storage_calgary["Calgary Status"] = "Not assessed - hydraulic routing solution invalid"
    if not criteria.suppress_empty_sections or not storage_calgary.empty:
        _add_df_table(doc, "Table 13 - Storage Unit Performance", storage_calgary, landscape=True, font_size=7.0)
    if criteria.suppress_empty_sections and control_table.empty and storage_table.empty:
        doc.add_paragraph("The model does not contain hydraulic controls or storage units; related result tables are not applicable.")

    doc.add_heading("4.5 HGL, Surcharge, and Flooding Assessment", level=2)
    _add_narrative_block(doc, narrative_sections, "hgl_surcharge")
    hgl_display = hgl_table.copy()
    if not hgl_display.empty and "Type" in hgl_display:
        mask = hgl_display["Type"].astype(str).str.lower().eq("outfall")
        for c in [f"Ground/Rim ({units.length})", f"Freeboard ({units.length})"]:
            if c in hgl_display:
                hgl_display[c] = hgl_display[c].astype(object)
                hgl_display.loc[mask, c] = "N/A"
        if "Status" in hgl_display: hgl_display.loc[mask, "Status"] = "Boundary condition"
    _add_df_table(doc, "Table 15 - Summary of Node HGL / Surcharge Conditions", hgl_display, landscape=True)
    doc.add_heading("4.6 Outfalls and Boundary Conditions", level=2)
    _add_narrative_block(doc, narrative_sections, "outfalls")
    _add_df_table(doc, "Table 16 - Minor System Boundary Conditions - Outflows", outfalls[[c for c in outfall_cols if c in outfalls.columns]])
    if narrative_sections and str(narrative_sections.get("private_sites", "")).strip():
        doc.add_heading("4.7 Private-Site Discharge and Storage", level=2)
        _add_narrative_block(doc, narrative_sections, "private_sites")
    if narrative_sections and str(narrative_sections.get("water_quality", "")).strip():
        doc.add_heading("4.8 Water Quality and BMPs", level=2)
        _add_narrative_block(doc, narrative_sections, "water_quality")
    if narrative_sections and str(narrative_sections.get("model_documentation", "")).strip():
        doc.add_heading("4.9 Model Input/Output Documentation", level=2)
        _add_narrative_block(doc, narrative_sections, "model_documentation")

    doc.add_heading("4.9 Continuity and Result Reconciliation", level=2)
    doc.add_paragraph(
        "Continuity errors are engine-reported values with sign preserved; thresholds are applied to their "
        "absolute magnitudes. Reconciliation compares the API-derived result tables against the engine's own "
        ".rpt summaries; for reconciliation-flagged links the .rpt value governs screening.")
    from screening_logic import continuity_disclosure as _cd
    for line in _cd(simulation_metadata, review_pct=criteria.continuity_review,
                    warning_pct=criteria.continuity_warning,
                    has_pollutants=bool(simulation_metadata.get("has_pollutants"))):
        doc.add_paragraph(line, style="List Bullet")
    if reconciliation and reconciliation.get("summary"):
        _rs = reconciliation["summary"]
        doc.add_paragraph(
            f"Reconciliation verdict: {_rs.get('verdict', 'Not performed')} "
            f"({_rs.get('ok', 0)} of {_rs.get('links_checked', 0)} links agree; "
            f"{int(_rs.get('review', 0)) + int(_rs.get('discrepancy', 0)) + int(_rs.get('unmatched', 0))} flagged). "
            "Reconciliation is not claimed to be fully clean while any link remains flagged.")
    if not effective_table.empty:
        _flagged_eff = effective_table[effective_table["Classification Changed by Reconciliation"] != "n/a - not flagged"]
        if not _flagged_eff.empty:
            _add_df_table(doc, "Table 13A - Reconciliation-Flagged Links: Evidence Precedence",
                          _flagged_eff, landscape=True, font_size=7.5)
            for _, _er in _flagged_eff.iterrows():
                changed = str(_er["Classification Changed by Reconciliation"]) == "Yes"
                doc.add_paragraph(
                    f"Link {_er['Link ID']}: worker {_er['Worker Peak Velocity (m/s)']} m/s vs engine .rpt "
                    f"{_er['RPT Peak Velocity (m/s)']} m/s; the .rpt value governs screening. The discrepancy "
                    f"{'CHANGES' if changed else 'does not change'} the screening classification "
                    f"({_er['Screening Classification']}).", style="List Bullet")

    doc.add_heading("4.10 Missing-Information Register", level=2)
    doc.add_paragraph(
        "Evidence items the model and session cannot supply. Any criterion depending on a listed item is "
        "reported as Not assessed rather than Pass.")
    _missing = missing_information_register(asdict(metadata), criteria_table if criteria.calgary_enabled else None, checklist_table)
    _add_df_table(doc, "Table 14A - Missing-Information Register", _missing, font_size=8.0)

    doc.add_heading("5.0 SUMMARY OF FINDINGS, CONCLUSIONS, AND RECOMMENDATIONS", level=1)
    _pda = (preliminary_review_artifacts or {}).get("findings")
    if _pda is not None:
        _pda_df = _pda if isinstance(_pda, pd.DataFrame) else pd.DataFrame(list(_pda))
        if not _pda_df.empty and "severity" in _pda_df.columns:
            _order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
            _pda_df = _pda_df.assign(_rank=_pda_df["severity"].astype(str).str.lower().map(_order).fillna(9)).sort_values("_rank")
            doc.add_heading("5.0a Prioritized Actions", level=2)
            for _, _f in _pda_df.head(10).iterrows():
                doc.add_paragraph(
                    f"[{_f.get('severity')}] {_f.get('finding_id')} — {str(_f.get('deterministic_basis', _f.get('finding_type', '')))[:220]} "
                    f"Action: {str(_f.get('recommended_action', 'Review.'))[:180]}", style="List Number")
    findings = _summary_findings(node_report, link_report, sub_report, simulation_metadata, units, options, criteria)
    doc.add_heading("5.1 Executive Summary", level=2)
    if not _add_narrative_block(doc, narrative_sections, "executive_summary"):
        for line in _executive_summary(findings, critical_nodes, critical_links, units, criteria):
            doc.add_paragraph(line)
    doc.add_heading("5.2 Detailed Findings", level=2)
    for finding in findings:
        doc.add_paragraph(finding, style="List Bullet")
    doc.add_paragraph(
        f"QA/QC screening criteria for junctions/dividers: depth ratio ≥ {criteria.node_depth_ratio:.2f}; rim clearance ≤ {criteria.minimum_freeboard:g} {units.length}; storage facilities are assessed separately by classification; "
        f"conduit depth ratio ≥ {criteria.conduit_depth_ratio:.2f}; velocity ≥ {criteria.velocity_threshold:g} {units.velocity}; "
        f"continuity review/warning thresholds = {criteria.continuity_review:g}%/{criteria.continuity_warning:g}%.",
        style=None,
    )
    doc.add_paragraph(
        "Peak-flow runoff coefficient is calculated from the Rational Method relationship and is a screening statistic; it is not necessarily equivalent to the volumetric event runoff coefficient.",
        style=None,
    )
    if not criteria.suppress_empty_sections or not critical_nodes.empty:
        _add_df_table(doc, "Table 17A - Critical Nodes Requiring Review", critical_nodes, font_size=8.0)
    if not criteria.suppress_empty_sections or not critical_links.empty:
        _add_df_table(doc, "Table 17B - Critical Conduits Requiring Review", critical_links, landscape=True, font_size=7.8)
    doc.add_heading("5.3 Outstanding Information and Actions", level=2)
    outstanding = checklist_table[~checklist_table["Status"].isin(["Complete", "Not applicable"])][["Item","Requirement","Status","Outstanding Action"]]
    _add_df_table(doc, "Table 18 - Outstanding SWMR Information and Actions", outstanding, landscape=True, font_size=7.2)
    _add_narrative_block(doc, narrative_sections, "outstanding_actions")
    doc.add_heading("5.4 Conclusions and Recommendations", level=2)
    _add_narrative_block(doc, narrative_sections, "conclusions")
    doc.add_paragraph(
        "This automatically generated report is a model-data population aid. It does not replace engineering judgment, "
        "municipal criteria verification, design-drawing review, boundary-condition confirmation, or professional authentication."
    )

    doc.add_page_break()
    doc.add_heading("APPENDIX A - MODEL DATA", level=1)
    link_geom, link_params, link_results = _split_link_appendix(conduit_table, units)
    _add_df_table(doc, "Table A-1 - Link Connectivity and Geometry", link_geom, landscape=True, font_size=7.0)
    _add_df_table(doc, "Table A-2 - Link Hydraulic Parameters", link_params, landscape=True, font_size=7.2)
    _add_df_table(doc, "Table A-3 - Link Simulation Results", link_results, landscape=True, font_size=7.2)
    if not criteria.suppress_empty_sections or not control_table.empty:
        _add_df_table(doc, "Table A-4 - Hydraulic Control Parameters and Results", control_table, landscape=True, font_size=7.0)
    if not criteria.suppress_empty_sections or not storage_calgary.empty:
        _add_df_table(doc, "Table A-5 - Storage Unit Parameters and Results", storage_calgary, landscape=True, font_size=6.8)
    _add_df_table(doc, "Table A-6 - Node Model and Result Summary", hgl_display, landscape=True, font_size=7.0)

    if criteria.calgary_enabled:
        doc.add_heading("APPENDIX B - CALGARY QA/QC", level=1)
        _add_df_table(doc, "Table B-1 - Project Criteria Register", criteria_table, landscape=True, font_size=7.2)
        if not overland_compliance.empty:
            _add_df_table(doc, "Table B-2 - Major-System Depth-Velocity Screening", overland_compliance, landscape=True, font_size=7.0)
        if not minor_capacity.empty:
            _add_df_table(doc, "Table B-3 - Minor-System Capacity Screening", minor_capacity, landscape=True, font_size=7.0)
        _add_df_table(doc, "Table B-4 - Calgary SWMR Completeness Register", checklist_table, landscape=True, font_size=6.8)
        _add_df_table(doc, "Table B-5 - Draft Readiness Summary", readiness_table, font_size=8.0)

    doc.add_heading("APPENDIX C - SUBCATCHMENT DATA", level=1)
    sub_model, sub_results = _split_subcatchment_appendix(sub_table, units)
    _add_df_table(doc, "Table B-1 - Subcatchment Model Parameters", sub_model, landscape=True, font_size=7.0)
    _add_df_table(doc, "Table B-2 - Subcatchment Simulation Results", sub_results, landscape=True, font_size=7.2)

    if model_listings and (model_listings.get("inp_text") or model_listings.get("rpt_text")):
        doc.add_heading("APPENDIX D - MODEL INPUT AND OUTPUT LISTINGS", level=1)
        doc.add_paragraph(
            "Fixed-width listings of the computer model input file and the engine output "
            "report are reproduced below (Calgary SWMR checklist item: attach computer model "
            "input and output files). The digital files in the accompanying audit package "
            "(model/ directory) are the authoritative machine-readable record.")
        if model_listings.get("inp_text"):
            _add_fixed_width_listing(
                doc, f"D.1 Model Input File ({model_listings.get('inp_name', 'model.inp')})",
                model_listings["inp_text"])
        if model_listings.get("rpt_text"):
            _add_fixed_width_listing(
                doc, f"D.2 Engine Output Report ({model_listings.get('rpt_name', 'model.rpt')})",
                model_listings["rpt_text"])

    llm_context = build_llm_report_context(
        metadata=asdict(metadata), criteria=calgary, findings=findings,
        tables={
            "design_event": event_table, "criteria_register": criteria_table,
            "minor_system": minor_capacity, "major_system": overland_compliance,
            "storage": storage_calgary, "critical_nodes": critical_nodes,
            "critical_conduits": critical_links, "outfalls": outfalls,
            "swmr_checklist": checklist_table, "draft_readiness": readiness_table,
            "scenario_comparison": scenario_comparison if scenario_comparison is not None else pd.DataFrame(),
        },
    )
    docx_buffer = io.BytesIO()
    _embed_attached_figures(doc, attached_figures)
    doc.save(docx_buffer); docx_bytes = docx_buffer.getvalue()
    base = _safe_name(metadata.project_name)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{base}_SWM_Report.docx", docx_bytes)
        zf.writestr("tables/node_summary.csv", node_report.to_csv(index=False) if node_summary is not None else "")
        zf.writestr("tables/link_summary.csv", link_report.to_csv(index=False) if link_summary is not None else "")
        zf.writestr("tables/subcatchment_summary.csv", sub_report.to_csv(index=False) if sub_summary is not None else "")
        zf.writestr("tables/catchment_model_results.csv", sub_table.to_csv(index=False))
        zf.writestr("tables/link_model_results.csv", link_table.to_csv(index=False))
        zf.writestr("tables/node_hgl_summary.csv", hgl_table.to_csv(index=False))
        zf.writestr("tables/hydraulic_controls.csv", control_table.to_csv(index=False))
        zf.writestr("tables/control_flow_reconciliation.csv", control_recon.to_csv(index=False))
        zf.writestr("tables/storage_unit_performance.csv", storage_table.to_csv(index=False))
        zf.writestr("tables/model_area_classification.csv", area_table.to_csv(index=False))
        zf.writestr("tables/critical_nodes.csv", critical_nodes.to_csv(index=False))
        zf.writestr("tables/critical_conduits.csv", critical_links.to_csv(index=False))
        zf.writestr("tables/design_event_summary.csv", event_table.to_csv(index=False))
        zf.writestr("tables/calgary_criteria_register.csv", criteria_table.to_csv(index=False))
        zf.writestr("tables/calgary_minor_system_capacity.csv", minor_capacity.to_csv(index=False))
        zf.writestr("tables/calgary_major_system_compliance.csv", overland_compliance.to_csv(index=False))
        zf.writestr("tables/depth_velocity_criterion.csv", curve_table.to_csv(index=False))
        zf.writestr("figures/figure_9_1_depth_velocity_criteria.png", depth_velocity_png)
        zf.writestr("figures/figure_3_1_model_schematic.png", schematic_png)
        zf.writestr("metadata/model_schematic_manifest.json", json.dumps(schematic_manifest, indent=2))
        zf.writestr("tables/calgary_storage_assessment.csv", storage_calgary.to_csv(index=False))
        zf.writestr("tables/calgary_swmr_completeness_register.csv", checklist_table.to_csv(index=False))
        zf.writestr("tables/swmr_draft_readiness.csv", readiness_table.to_csv(index=False))
        if scenario_comparison is not None and not scenario_comparison.empty:
            zf.writestr("tables/preliminary_scenario_comparison.csv", scenario_comparison.to_csv(index=False))
        if scenario_analysis:
            zf.writestr("narratives/preliminary_scenario_comparison_analysis.txt", str(scenario_analysis))
        if scenario_records:
            for record in scenario_records:
                sid = str((record.get("definition", {}) or {}).get("scenario_id", "scenario"))
                if sid == "BASE_MODEL":
                    continue
                for key, frame in _scenario_detail_tables(record, units).items():
                    if frame is not None and not frame.empty:
                        zf.writestr(f"tables/scenarios/{sid}_{key}.csv", frame.to_csv(index=False))
        zf.writestr("metadata/llm_report_context.json", json.dumps(llm_context, indent=2, default=str))
        if not effective_table.empty:
            zf.writestr("tables/velocity_screening_effective.csv", effective_table.to_csv(index=False))
        if model_identity:
            zf.writestr("metadata/model_identity.json", json.dumps(dict(model_identity), indent=2, default=str))
        if model_listings:
            if model_listings.get("inp_text"):
                zf.writestr(f"model/{_safe_name(model_listings.get('inp_name', 'model.inp'))}",
                            model_listings["inp_text"])
            if model_listings.get("execution_inp_text"):
                zf.writestr(f"model/{_safe_name(model_listings.get('execution_inp_name', 'execution_model.inp'))}",
                            model_listings["execution_inp_text"])
            if model_listings.get("rpt_text"):
                zf.writestr(f"model/{_safe_name(model_listings.get('rpt_name', 'model.rpt'))}",
                            model_listings["rpt_text"])
        if attached_figures:
            fig_manifest = []
            for n, fig in enumerate(attached_figures, start=1):
                fpath = Path(str(fig.get("path", "")))
                if fpath.exists():
                    zf.writestr(f"figures/{fpath.name}", fpath.read_bytes())
                fig_manifest.append({"n": n, "figure_id": fig.get("figure_id"),
                                     "file": fpath.name, "caption": fig.get("caption"),
                                     "section": fig.get("section"), "source": fig.get("source"),
                                     "note": "Client-attached illustrative figure; not a server-verified output."})
            zf.writestr("metadata/attached_figures.json", json.dumps(fig_manifest, indent=2, default=str))
        zf.writestr("metadata/approved_narrative_sections.json", json.dumps(dict(narrative_sections or {}), indent=2, ensure_ascii=False))
        zf.writestr("metadata/project_metadata.json", json.dumps(asdict(metadata), indent=2))
        zf.writestr("metadata/report_criteria.json", json.dumps(asdict(criteria), indent=2, default=str))
        zf.writestr("metadata/simulation_metadata.json", json.dumps(simulation_metadata, indent=2, default=str))
        zf.writestr("metadata/inp_sections.json", json.dumps(inp_sections, indent=2, default=str))
        if preliminary_review_artifacts:
            zf.writestr("preliminary_design/review_manifest.json", json.dumps(preliminary_review_artifacts.get("manifest", {}), indent=2, default=str))
            findings_df = preliminary_review_artifacts.get("findings")
            if isinstance(findings_df, pd.DataFrame):
                zf.writestr("preliminary_design/findings_register.csv", findings_df.to_csv(index=False))
            elif findings_df:
                zf.writestr("preliminary_design/findings_register.json", json.dumps(findings_df, indent=2, default=str))
            if preliminary_review_artifacts.get("ai_review"):
                zf.writestr("preliminary_design/ai_review.md", str(preliminary_review_artifacts.get("ai_review")))
            if preliminary_review_artifacts.get("reviewed_model"):
                zf.writestr("preliminary_design/reviewed_scenario_base.inp", str(preliminary_review_artifacts.get("reviewed_model")))
        if result_db_bytes:
            zf.writestr("database/swmm_complete_results.sqlite", result_db_bytes)

    return {
        "docx": docx_bytes,
        "zip": zip_buffer.getvalue(),
        "docx_name": f"{base}_SWM_Report.docx",
        "zip_name": f"{base}_SWM_Report_Package.zip",
        "catchment_table": sub_table,
        "link_table": link_table,
        "hgl_table": hgl_table,
        "control_table": control_table,
        "storage_table": storage_calgary,
        "area_table": area_table,
        "calgary_criteria_table": criteria_table,
        "calgary_minor_capacity": minor_capacity,
        "calgary_overland_compliance": overland_compliance,
        "swmr_checklist": checklist_table, "draft_readiness": readiness_table,
        "llm_context": llm_context,
        "narrative_sections": dict(narrative_sections or {}),
        "scenario_comparison": scenario_comparison if scenario_comparison is not None else pd.DataFrame(),
    }
